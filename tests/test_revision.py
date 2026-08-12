#!/usr/bin/env python3
"""Tests de revision.py: checks de formato puros, reglas, comparacion y
gestion de errores (sin modelos ni LibreOffice: se simulan).

Ejecutar: python3 -m unittest discover -s tests -v
"""

import json
import os
import tempfile
import unittest
from unittest import mock

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from revision import (CHECKS, CHECKS_DOCX, CHECKS_PDF, REGLAS_DEFAULT,
                      cargar_reglas, comparar_planillas, rango_datos,
                      revisar_documento, revisar_planilla, vision_360,
                      _parsear_rubrica, _indice_a_letra)


def _soffice_disponible() -> bool:
    import shutil

    return bool(shutil.which("soffice") or shutil.which("libreoffice"))


def _hoja_con(contenido):
    """Hoja en memoria a partir de una lista de filas."""
    wb = Workbook()
    hoja = wb.active
    hoja.title = "Datos"
    for fila in contenido:
        hoja.append(fila)
    return hoja


def _ctx(hoja):
    return {"rango": rango_datos(hoja), "max_hallazgos": 100}


def _borde(celda):
    b = Side(style="thin", color="000000")
    celda.border = Border(left=b, right=b, top=b, bottom=b)


def _guardar(hoja):
    """Guarda el libro de la hoja en un archivo temporal y devuelve la ruta."""
    fd, ruta = tempfile.mkstemp(suffix=".xlsx")
    os.close(fd)
    hoja.parent.save(ruta)
    return ruta


class TestReglas(unittest.TestCase):

    def test_defaults_sin_archivo(self):
        reglas, errores = cargar_reglas()
        self.assertEqual(errores, [])
        todos = set(CHECKS) | set(CHECKS_DOCX) | set(CHECKS_PDF)
        self.assertEqual(set(reglas), todos)
        self.assertTrue(reglas["encabezados"]["activo"])
        self.assertTrue(reglas["docx_margenes"]["activo"])
        self.assertTrue(reglas["pdf_rotacion"]["activo"])

    def test_override_y_errores(self):
        fd, ruta = tempfile.mkstemp(suffix=".json")
        os.close(fd)
        with open(ruta, "w", encoding="utf-8") as f:
            json.dump({"encabezados": {"activo": False, "negrita_requerida": False},
                       "check_inventado": {"activo": True}}, f)
        reglas, errores = cargar_reglas(ruta)
        self.assertIn("check desconocido: 'check_inventado'", errores)
        self.assertFalse(reglas["encabezados"]["activo"])
        self.assertFalse(reglas["encabezados"]["negrita_requerida"])
        os.unlink(ruta)

    def test_severidad_invalida(self):
        fd, ruta = tempfile.mkstemp(suffix=".json")
        os.close(fd)
        with open(ruta, "w", encoding="utf-8") as f:
            json.dump({"bordes": {"severidad": "critica"}}, f)
        reglas, errores = cargar_reglas(ruta)
        self.assertTrue(any("severidad invalida" in e for e in errores))
        self.assertEqual(reglas["bordes"]["severidad"], "warning")
        os.unlink(ruta)

    def test_archivo_inexistente(self):
        reglas, errores = cargar_reglas("/no/existe/reglas.json")
        self.assertEqual(reglas, REGLAS_DEFAULT)
        self.assertEqual(len(errores), 1)


class TestChecks(unittest.TestCase):

    def test_encabezados_sin_negrita(self):
        hoja = _hoja_con([["A", "B"], [1, 2]])
        hallazgos = CHECKS["encabezados"](hoja, REGLAS_DEFAULT["encabezados"], _ctx(hoja))
        self.assertEqual(len(hallazgos), 2)
        self.assertEqual(hallazgos[0]["severidad"], "error")

    def test_encabezados_ok(self):
        hoja = _hoja_con([["A", "B"], [1, 2]])
        for c in (hoja["A1"], hoja["B1"]):
            c.font = Font(bold=True)
        hallazgos = CHECKS["encabezados"](hoja, REGLAS_DEFAULT["encabezados"], _ctx(hoja))
        self.assertEqual(hallazgos, [])

    def test_columna_sin_nombre_solo_si_tiene_datos(self):
        # B tiene datos pero sin nombre de encabezado -> reporta
        hoja = _hoja_con([["A", None, "C"], [1, 2, 3]])
        hallazgos = CHECKS["encabezados"](hoja, REGLAS_DEFAULT["encabezados"], _ctx(hoja))
        sin_nombre = [h["celda"] for h in hallazgos
                      if "sin nombre" in h["mensaje"]]
        self.assertEqual(sin_nombre, ["B1"])
        # B sin datos y D sin datos: ninguno reporta (falso positivo evitado)
        hoja2 = _hoja_con([["A", None, "C"], [1, None, 3]])
        hoja2["D1"] = "D"
        hallazgos2 = CHECKS["encabezados"](hoja2, REGLAS_DEFAULT["encabezados"], _ctx(hoja2))
        sin_nombre2 = [h["celda"] for h in hallazgos2
                       if "sin nombre" in h["mensaje"]]
        self.assertEqual(sin_nombre2, [])

    def test_bordes(self):
        hoja = _hoja_con([["A", "B"], [1, 2], [3, 4]])
        _borde(hoja["A1"]); _borde(hoja["B1"]); _borde(hoja["A2"]); _borde(hoja["B2"])
        hallazgos = CHECKS["bordes"](hoja, REGLAS_DEFAULT["bordes"], _ctx(hoja))
        self.assertEqual([h["celda"] for h in hallazgos], ["A3", "B3"])

    def test_alineacion_numerico_izquierda(self):
        hoja = _hoja_con([["A", "B"], [1, "texto"]])
        hoja["A2"].alignment = Alignment(horizontal="left")
        hoja["B2"].alignment = Alignment(horizontal="right")
        hallazgos = CHECKS["alineacion"](hoja, REGLAS_DEFAULT["alineacion"], _ctx(hoja))
        celdas = {h["celda"]: h for h in hallazgos}
        self.assertIn("A2", celdas)
        self.assertIn("B2", celdas)

    def test_anchos(self):
        hoja = _hoja_con([["A"], [1]])
        hoja.column_dimensions["A"].width = 3
        hallazgos = CHECKS["anchos"](hoja, REGLAS_DEFAULT["anchos"], _ctx(hoja))
        self.assertEqual(len(hallazgos), 1)
        self.assertIn("minimo", hallazgos[0]["mensaje"])

    def test_formato_numero_texto(self):
        hoja = _hoja_con([["A"], ["120"], [120.5]])
        hoja["A3"].number_format = "General"
        hallazgos = CHECKS["formato_numero"](hoja, REGLAS_DEFAULT["formato_numero"], _ctx(hoja))
        self.assertEqual(len(hallazgos), 2)  # '120' como texto + 120.5 en General

    def test_formato_numero_entero_general_ok(self):
        hoja = _hoja_con([["A"], [2024], [7]])
        hallazgos = CHECKS["formato_numero"](hoja, REGLAS_DEFAULT["formato_numero"], _ctx(hoja))
        self.assertEqual(hallazgos, [])

    def test_filtros_requerido(self):
        hoja = _hoja_con([["A", "B"], [1, 2]])
        hallazgos = CHECKS["filtros"](hoja, REGLAS_DEFAULT["filtros"], _ctx(hoja))
        self.assertEqual(len(hallazgos), 1)
        hoja.auto_filter.ref = "A1:B2"
        hallazgos = CHECKS["filtros"](hoja, REGLAS_DEFAULT["filtros"], _ctx(hoja))
        self.assertEqual(hallazgos, [])

    def test_celdas_vacias(self):
        hoja = _hoja_con([["A", "B", "C"], [1, None, 3]])
        hallazgos = CHECKS["celdas_vacias"](hoja, REGLAS_DEFAULT["celdas_vacias"], _ctx(hoja))
        self.assertEqual([h["celda"] for h in hallazgos], ["B2"])

    def test_celdas_mezcladas(self):
        hoja = _hoja_con([["A"], [1]])
        hoja.merge_cells("A2:B2")
        hallazgos = CHECKS["celdas_mezcladas"](hoja, REGLAS_DEFAULT["celdas_mezcladas"], _ctx(hoja))
        self.assertEqual(len(hallazgos), 1)
        self.assertIn("fusionadas", hallazgos[0]["mensaje"])

    def test_ocultas(self):
        hoja = _hoja_con([["A", "B"], [1, 2]])
        hoja.row_dimensions[2].hidden = True
        hoja.column_dimensions["B"].hidden = True
        hallazgos = (CHECKS["filas_ocultas"](hoja, REGLAS_DEFAULT["filas_ocultas"], _ctx(hoja))
                     + CHECKS["columnas_ocultas"](hoja, REGLAS_DEFAULT["columnas_ocultas"], _ctx(hoja)))
        self.assertEqual(len(hallazgos), 2)

    def test_errores_formula(self):
        hoja = _hoja_con([["A"], ["#DIV/0!"], [1]])
        hallazgos = CHECKS["errores_formula"](hoja, REGLAS_DEFAULT["errores_formula"], _ctx(hoja))
        self.assertEqual([h["celda"] for h in hallazgos], ["A2"])

    def test_duplicados_encabezado(self):
        hoja = _hoja_con([["A", "A", "B"], [1, 2, 3]])
        hallazgos = CHECKS["duplicados_encabezado"](
            hoja, REGLAS_DEFAULT["duplicados_encabezado"], _ctx(hoja))
        self.assertEqual(len(hallazgos), 1)
        self.assertEqual(hallazgos[0]["detalle"]["nombre"], "A")

    def test_texto_desbordado(self):
        hoja = _hoja_con([["A", "B"], ["texto muy largo que no cabe", "ocupado"]])
        hoja.column_dimensions["A"].width = 6
        hallazgos = CHECKS["texto_desbordado"](hoja, REGLAS_DEFAULT["texto_desbordado"], _ctx(hoja))
        self.assertEqual([h["celda"] for h in hallazgos], ["A2"])
        hoja["B2"].value = None  # vecino libre: el texto puede desbordarse
        hallazgos = CHECKS["texto_desbordado"](hoja, REGLAS_DEFAULT["texto_desbordado"], _ctx(hoja))
        self.assertEqual(hallazgos, [])

    def test_estilos_inconsistentes(self):
        hoja = _hoja_con([["A", "B"], [1.5, 2.5], [3.5, 4.5]])
        hoja["B2"].number_format = "0.00"
        hallazgos = CHECKS["estilos_inconsistentes"](
            hoja, REGLAS_DEFAULT["estilos_inconsistentes"], _ctx(hoja))
        self.assertEqual([h["celda"] for h in hallazgos], ["columna B"])

    def test_islas_datos(self):
        hoja = _hoja_con([["A"], [1], [None], [None], [None], [None], [2]])
        hallazgos = CHECKS["islas_datos"](hoja, REGLAS_DEFAULT["islas_datos"], _ctx(hoja))
        self.assertEqual(len(hallazgos), 1)
        self.assertIn("hueco", hallazgos[0]["mensaje"])

    def test_proteccion_requerida(self):
        hoja = _hoja_con([["A"], [1]])
        cfg = dict(REGLAS_DEFAULT["proteccion"], requerida=True)
        hallazgos = CHECKS["proteccion"](hoja, cfg, _ctx(hoja))
        self.assertEqual(len(hallazgos), 1)
        hoja.protection.sheet = True
        hallazgos = CHECKS["proteccion"](hoja, cfg, _ctx(hoja))
        self.assertEqual(hallazgos, [])


class TestRevisionCompleta(unittest.TestCase):

    def _buena(self):
        hoja = _hoja_con([["A", "B"], [1.5, "x"], [2.5, "y"]])
        for c in (hoja["A1"], hoja["B1"]):
            c.font = Font(bold=True)
            _borde(c)
        _borde(hoja["A2"]); _borde(hoja["B2"]); _borde(hoja["A3"]); _borde(hoja["B3"])
        hoja["A2"].number_format = "0.00"
        hoja["A3"].number_format = "0.00"
        hoja.auto_filter.ref = "A1:B3"
        hoja.column_dimensions["A"].width = 12
        hoja.column_dimensions["B"].width = 12
        return hoja

    def test_planilla_buena_sin_hallazgos(self):
        ruta = _guardar(self._buena())
        resultado = revisar_planilla(ruta)
        os.unlink(ruta)
        self.assertTrue(resultado["ok"])
        self.assertEqual(resultado["resumen"], {"error": 0, "warning": 0, "info": 0})
        self.assertEqual(resultado["hojas"], ["Datos"])

    def test_planilla_mala_detecta(self):
        hoja = self._buena()
        hoja["A1"].font = Font(bold=False)  # quitar negrita del encabezado
        hoja.auto_filter.ref = None         # quitar filtro
        hoja["A4"] = "dato suelto"          # isla
        ruta = _guardar(hoja)
        resultado = revisar_planilla(ruta)
        os.unlink(ruta)
        reglas_detectadas = {h["regla"] for h in resultado["hallazgos"]}
        self.assertIn("encabezados", reglas_detectadas)
        self.assertIn("filtros", reglas_detectadas)
        self.assertGreater(resultado["resumen"]["error"], 0)

    def test_hoja_inexistente(self):
        ruta = _guardar(self._buena())
        with self.assertRaises(ValueError):
            revisar_planilla(ruta, hoja_solo="NoExiste")
        os.unlink(ruta)

    def test_check_con_error_interno_no_mata_la_revision(self):
        ruta = _guardar(self._buena())
        with mock.patch.dict("revision.CHECKS",
                             {"encabezados": lambda h, c, x: (_ for _ in ()).throw(RuntimeError("boom"))}):
            resultado = revisar_planilla(ruta)
        os.unlink(ruta)
        self.assertTrue(resultado["ok"])
        self.assertEqual(resultado["hallazgos"][0]["regla"], "encabezados")
        self.assertIn("error interno", resultado["hallazgos"][0]["mensaje"])


class TestComparacion(unittest.TestCase):

    def test_diferencias(self):
        a = _hoja_con([["C", "V"], ["x", 1.0]])
        b = _hoja_con([["C", "V2"], ["x", 2.0], ["y", 3.0]])
        ruta_a, ruta_b = _guardar(a), _guardar(b)
        resultado = comparar_planillas(ruta_a, ruta_b)
        os.unlink(ruta_a); os.unlink(ruta_b)
        tipos = {d["tipo"] for d in resultado["diferencias"]}
        self.assertIn("encabezado", tipos)
        self.assertIn("valor", tipos)
        self.assertIn("dimension", tipos)
        self.assertEqual(resultado["total_diferencias"], 5)

    def test_hojas_distintas(self):
        a = _hoja_con([["A"], [1]])
        b = _hoja_con([["A"], [1]])
        b.parent.create_sheet("Extra")
        ruta_a, ruta_b = _guardar(a), _guardar(b)
        resultado = comparar_planillas(ruta_a, ruta_b)
        os.unlink(ruta_a); os.unlink(ruta_b)
        self.assertIn("hoja", {d["tipo"] for d in resultado["diferencias"]})


class TestChecksDocx(unittest.TestCase):
    """Checks de documento (python-docx, en memoria)."""

    def _doc(self):
        from docx import Document

        return Document()

    def test_titulos_estilos(self):
        doc = self._doc()
        doc.add_heading("Titulo real", level=1)
        p = doc.add_paragraph()
        r = p.add_run("Titulo manual")
        r.bold = True
        hallazgos = CHECKS_DOCX["docx_titulos_estilos"](
            doc, REGLAS_DEFAULT["docx_titulos_estilos"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertIn("manual", hallazgos[0]["mensaje"])

    def test_fuentes(self):
        doc = self._doc()
        p = doc.add_paragraph()
        p.add_run("A").font.name = "Calibri"
        p.add_run("B").font.name = "Times New Roman"
        p.add_run("C").font.name = "Arial"
        hallazgos = CHECKS_DOCX["docx_fuentes"](
            doc, REGLAS_DEFAULT["docx_fuentes"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertIn("3 fuentes", hallazgos[0]["mensaje"])

    def test_margenes(self):
        from docx.shared import Cm

        doc = self._doc()
        doc.sections[0].left_margin = Cm(0.3)
        hallazgos = CHECKS_DOCX["docx_margenes"](
            doc, REGLAS_DEFAULT["docx_margenes"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertIn("izquierdo", hallazgos[0]["mensaje"])

    def test_numeracion_manual(self):
        doc = self._doc()
        doc.add_paragraph("1. Paso uno")
        doc.add_paragraph("Paso sin numero")
        doc.add_paragraph("2. Paso dos")
        hallazgos = CHECKS_DOCX["docx_numeracion_manual"](
            doc, REGLAS_DEFAULT["docx_numeracion_manual"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 2)

    def test_parrafos_vacios(self):
        doc = self._doc()
        doc.add_paragraph("texto")
        for _ in range(4):
            doc.add_paragraph("")
        doc.add_paragraph("mas texto")
        hallazgos = CHECKS_DOCX["docx_parrafos_vacios"](
            doc, REGLAS_DEFAULT["docx_parrafos_vacios"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertEqual(hallazgos[0]["detalle"]["rachas"], [4])

    def test_tablas_sin_estilo(self):
        doc = self._doc()
        tabla = doc.add_table(rows=1, cols=2)
        # el default de python-docx es 'Normal Table' (sin bordes visibles)
        hallazgos = CHECKS_DOCX["docx_tablas_sin_estilo"](
            doc, REGLAS_DEFAULT["docx_tablas_sin_estilo"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        tabla.style = "Table Grid"
        hallazgos = CHECKS_DOCX["docx_tablas_sin_estilo"](
            doc, REGLAS_DEFAULT["docx_tablas_sin_estilo"], {"max_hallazgos": 100})
        self.assertEqual(hallazgos, [])

    def test_encabezados_pie_informativo(self):
        doc = self._doc()
        hallazgos = CHECKS_DOCX["docx_encabezados_pie"](
            doc, REGLAS_DEFAULT["docx_encabezados_pie"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertEqual(hallazgos[0]["severidad"], "info")

    def test_imagenes(self):
        doc = self._doc()
        hallazgos = CHECKS_DOCX["docx_imagenes"](
            doc, REGLAS_DEFAULT["docx_imagenes"], {"max_hallazgos": 100})
        self.assertEqual(hallazgos, [])  # sin imagenes: no reporta


class TestChecksPdf(unittest.TestCase):
    """Checks de PDF (datos de paginas simulados, sin archivo)."""

    def _paginas(self, *textos, rotaciones=None, tamanos=None):
        n = len(textos)
        rotaciones = rotaciones or [0] * n
        tamanos = tamanos or [(612, 792)] * n
        return [{"indice": i, "texto": t, "rotacion": rotaciones[i],
                 "ancho": tamanos[i][0], "alto": tamanos[i][1]}
                for i, t in enumerate(textos)]

    def test_paginas_vacias(self):
        paginas = self._paginas("hola", "   ", "mundo")
        hallazgos = CHECKS_PDF["pdf_paginas_vacias"](
            paginas, REGLAS_DEFAULT["pdf_paginas_vacias"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertIn("1 pagina", hallazgos[0]["mensaje"])
        self.assertEqual(hallazgos[0]["detalle"]["paginas"], [2])

    def test_paginas_escasas(self):
        paginas = self._paginas("hola", "x" * 5, "y" * 500)
        hallazgos = CHECKS_PDF["pdf_paginas_escasas"](
            paginas, REGLAS_DEFAULT["pdf_paginas_escasas"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertEqual(hallazgos[0]["detalle"]["paginas"],
                         [{"pagina": 1, "caracteres": 4},
                          {"pagina": 2, "caracteres": 5}])

    def test_sin_capa_texto(self):
        paginas = self._paginas("", "  ")
        hallazgos = CHECKS_PDF["pdf_sin_capa_texto"](
            paginas, REGLAS_DEFAULT["pdf_sin_capa_texto"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertIn("escaneado", hallazgos[0]["mensaje"])
        paginas = self._paginas("con texto", "mas")
        hallazgos = CHECKS_PDF["pdf_sin_capa_texto"](
            paginas, REGLAS_DEFAULT["pdf_sin_capa_texto"], {"max_hallazgos": 100})
        self.assertEqual(hallazgos, [])

    def test_rotacion(self):
        paginas = self._paginas("a", "b", "c", rotaciones=[0, 90, 0])
        hallazgos = CHECKS_PDF["pdf_rotacion"](
            paginas, REGLAS_DEFAULT["pdf_rotacion"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)
        self.assertEqual(hallazgos[0]["detalle"]["paginas"], [2])

    def test_tamano_paginas(self):
        paginas = self._paginas("a", "b",
                                tamanos=[(612, 792), (612, 792)])
        hallazgos = CHECKS_PDF["pdf_tamano_paginas"](
            paginas, REGLAS_DEFAULT["pdf_tamano_paginas"], {"max_hallazgos": 100})
        self.assertEqual(hallazgos, [])
        paginas = self._paginas("a", "b",
                                tamanos=[(612, 792), (792, 612)])
        hallazgos = CHECKS_PDF["pdf_tamano_paginas"](
            paginas, REGLAS_DEFAULT["pdf_tamano_paginas"], {"max_hallazgos": 100})
        self.assertEqual(len(hallazgos), 1)


class TestDocumento(unittest.TestCase):

    def test_archivo_inexistente(self):
        resultado = revisar_documento("/no/existe.xlsx")
        self.assertFalse(resultado["ok"])
        self.assertIn("no existe", resultado["error"])

    def test_formato_no_soportado(self):
        resultado = revisar_documento("ejemplos/grafico_demo.png")
        self.assertFalse(resultado["ok"])
        self.assertIn("no soportado", resultado["error"])

    def test_xlsx_ok(self):
        hoja = _hoja_con([["A"], [1]])
        ruta = _guardar(hoja)
        resultado = revisar_documento(ruta)
        os.unlink(ruta)
        self.assertTrue(resultado["ok"])
        self.assertEqual(resultado["formato"], "xlsx")

    def test_hoja_solo_restriccion_formato(self):
        # --hoja solo aplica a planillas, no a docx/pdf
        resultado = revisar_documento("ejemplos/planillas/documento.pdf",
                                      hoja_solo="Ventas")
        self.assertFalse(resultado["ok"])
        self.assertIn("--hoja solo aplica", resultado["error"])

    def test_comparar_restriccion_formato(self):
        resultado = revisar_documento("ejemplos/planillas/documento.pdf",
                                      comparar="ejemplos/planillas/v1.xlsx")
        self.assertFalse(resultado["ok"])
        self.assertIn("--comparar solo aplica", resultado["error"])

    @unittest.skipUnless(_soffice_disponible(), "requiere soffice/libreoffice")
    def test_ods_con_soffice(self):
        resultado = revisar_documento("ejemplos/planillas/correcta.ods")
        self.assertTrue(resultado["ok"])
        self.assertEqual(resultado["formato"], "ods")
        self.assertIn("transformacion", resultado)

    @unittest.skipUnless(_soffice_disponible(), "requiere soffice/libreoffice")
    def test_comparar_ods_normaliza(self):
        # openpyxl no lee ods: --comparar debe normalizar ambos lados a xlsx
        resultado = revisar_documento("ejemplos/planillas/correcta.ods",
                                      comparar="ejemplos/planillas/v1.xlsx")
        self.assertTrue(resultado["ok"])
        self.assertIn("comparacion", resultado)
        self.assertGreaterEqual(resultado["comparacion"]["total_diferencias"], 1)

    @unittest.skipUnless(_soffice_disponible(), "requiere soffice/libreoffice")
    def test_docx_con_soffice_no_necesario(self):
        resultado = revisar_documento("ejemplos/planillas/documento_correcto.docx")
        self.assertTrue(resultado["ok"])
        self.assertEqual(resultado["formato"], "docx")
        self.assertEqual(resultado["resumen"]["error"], 0)

    def test_pdf(self):
        resultado = revisar_documento("ejemplos/planillas/documento.pdf")
        self.assertTrue(resultado["ok"])
        self.assertEqual(resultado["formato"], "pdf")
        self.assertGreaterEqual(resultado["estructura"]["paginas"], 1)


class TestRubricaVLM(unittest.TestCase):

    def test_parseo_rubrica(self):
        texto = ("legibilidad: 8/10 | buen contraste\n"
                 "coherencia de estilo: 7/10 | colores consistentes\n"
                 "formato de datos: 5/10 | decimales inconsistentes\n"
                 "basura sin formato")
        notas, comentarios, no_conformes = _parsear_rubrica(texto)
        self.assertEqual(notas["legibilidad"], 8)
        self.assertEqual(notas["coherencia"], 7)
        self.assertEqual(notas["formato de datos"], 5)
        self.assertEqual(len(notas), 3)
        self.assertEqual(no_conformes, 0)
        self.assertIn("buen contraste", comentarios["legibilidad"])

    def test_parseo_nota_primero(self):
        # docbee responde a veces '10/10: Legibilidad | comentario'
        texto = ("10/10: Legibilidad excelente | contraste bueno\n"
                 "9/10: Coherencia de estilo | colores adecuados")
        notas, comentarios, no_conformes = _parsear_rubrica(texto)
        self.assertEqual(notas["legibilidad"], 10)
        self.assertEqual(notas["coherencia"], 9)
        self.assertEqual(no_conformes, 0)

    def test_parseo_dimensiones_inventadas_no_conformes(self):
        # el VLM repite dimensiones fuera de rubrica: se cuentan y se ignoran
        texto = ("10/10: Legibilidad excelente | ok\n"
                 "9/10: Utilidad de la planilla excelente | repetido\n"
                 "8/10: Diseño de la planilla excelente | repetido")
        notas, _, no_conformes = _parsear_rubrica(texto)
        self.assertEqual(notas, {"legibilidad": 10})
        self.assertEqual(no_conformes, 2)

    def test_parseo_sin_notas(self):
        notas, _, no_conformes = _parsear_rubrica("el diseno es aceptable")
        self.assertEqual(notas, {})
        self.assertEqual(no_conformes, 0)


class TestVision(unittest.TestCase):

    def test_motor_invalido(self):
        with mock.patch("revision.render_libro_a_pngs") as render:
            resultado = vision_360("x.xlsx", motor="inventado")
        render.assert_not_called()
        self.assertFalse(resultado["ok"])
        self.assertIn("motor invalido", resultado["error"])

    def test_ollama_con_render(self):
        # bateria_360 no importable (None en sys.modules -> ImportError)
        with mock.patch("revision.render_libro_a_pngs", return_value=["pag1.png"]), \
                mock.patch.dict("sys.modules", {"bateria_360": None}):
            resultado = vision_360("x.xlsx", motor="ollama")
        self.assertFalse(resultado["ok"])
        self.assertIn("bateria_360", resultado["error"])

    def test_ollama_completo(self):
        bateria = mock.MagicMock()
        bateria.run_ollama.return_value = {
            "ok": True,
            "texto": "legibilidad: 9/10 | excelente\ncoherencia de estilo: 8/10 | ok",
            "total_s": 12.3,
        }
        with mock.patch("revision.render_libro_a_pngs", return_value=["pag1.png"]), \
                mock.patch.dict("sys.modules", {"bateria_360": bateria}):
            resultado = vision_360("x.xlsx", motor="ollama")
        self.assertTrue(resultado["ok"])
        self.assertEqual(resultado["resultados"][0]["notas"]["legibilidad"], 9)

    def test_render_soffice_ausente(self):
        with mock.patch("shutil.which", return_value=None):
            from revision import render_libro_a_pngs
            with self.assertRaises(RuntimeError):
                render_libro_a_pngs("x.xlsx")


class TestUtilidades(unittest.TestCase):

    def test_indice_a_letra(self):
        self.assertEqual(_indice_a_letra(1), "A")
        self.assertEqual(_indice_a_letra(27), "AA")

    def test_rango_vacio(self):
        hoja = _hoja_con([])
        self.assertEqual(rango_datos(hoja), (0, 0, 0, 0))


if __name__ == "__main__":
    unittest.main()
