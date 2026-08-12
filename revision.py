#!/usr/bin/env python3
"""Revision de formato y presentacion de planillas y documentos.

Dos capas complementarias, coherentes con el resto del proyecto:

  1. Analisis determinista por formato, configurables por JSON con defaults
     profesionales:
     - xlsx/xlsm (openpyxl): encabezados, bordes, alineacion, anchos,
       formato numerico, filtros, celdas vacias/mezcladas, filas y
       columnas ocultas, errores de formula, duplicados de encabezado,
       texto desbordado, estilos inconsistentes, islas de datos,
       proteccion; comparacion entre versiones (--comparar).
     - ods: normalizacion a xlsx via LibreOffice headless con verificacion
       de integridad de la conversion (odfpy opcional).
     - docx (python-docx): titulos con estilos, fuentes, margenes,
       numeracion manual, parrafos vacios, tablas sin bordes,
       encabezado/pie, imagenes.
     - pdf (pypdfium2): paginas vacias/escasas, sin capa de texto,
       rotacion, tamanos de pagina.

  2. Vision IA 360 (opt-in, --vision): render de cada pagina a imagen
     (PDF nativo con pypdfium2; el resto LibreOffice -> PDF -> pypdfium2 ->
     PNG) y evaluacion de diseno/presentacion por un VLM local (docbee /
     ollama), reutilizando los motores validados de scripts/bateria_360.py.

Uso:
  python3 revision.py planilla.xlsx [--reglas reglas.json] [--comparar otra]
      [--vision docbee|ollama] [--modelo gemma3:4b] [--host 127.0.0.1]
      [--device cuda] [--hoja Nombre] [--max-hallazgos N]

Cada hallazgo: {"regla", "severidad": error|warning|info, "hoja", "celda",
"mensaje", "detalle"}.
"""

import argparse
import json
import os
import re
import subprocess
import sys
import tempfile
import time

# openpyxl solo se usa en el analisis xlsx: import perezoso dentro de las
# funciones para que el modulo sea importable sin dependencias pesadas
# (mismo patron que vision.py con paddleocr).

# ---------------------------------------------------------------------------
# Reglas y defaults profesionales
# ---------------------------------------------------------------------------

SEVERIDADES = ("error", "warning", "info")

# Cada check: {activo, severidad, parametros}. Los parametros se pueden
# sobreescribir en el JSON de reglas; el archivo solo necesita incluir lo
# que cambia respecto a este default.
REGLAS_DEFAULT = {
    "encabezados": {
        "activo": True,
        "severidad": "error",
        "fila_encabezado": 1,
        "negrita_requerida": True,
        "color_fondo_requerido": False,
        "bordes_requeridos": False,
    },
    "bordes": {
        "activo": True,
        "severidad": "warning",
        "solo_exterior": False,
    },
    "alineacion": {
        "activo": True,
        "severidad": "warning",
    },
    "anchos": {
        "activo": True,
        "severidad": "warning",
        "minimo": 8,
        "maximo": 60,
    },
    "formato_numero": {
        "activo": True,
        "severidad": "warning",
    },
    "filtros": {
        "activo": True,
        "severidad": "warning",
        "requerido": True,
    },
    "celdas_vacias": {
        "activo": True,
        "severidad": "warning",
    },
    "celdas_mezcladas": {
        "activo": True,
        "severidad": "info",
    },
    "filas_ocultas": {
        "activo": True,
        "severidad": "error",
    },
    "columnas_ocultas": {
        "activo": True,
        "severidad": "error",
    },
    "errores_formula": {
        "activo": True,
        "severidad": "error",
    },
    "duplicados_encabezado": {
        "activo": True,
        "severidad": "error",
        "fila_encabezado": 1,
    },
    "texto_desbordado": {
        "activo": True,
        "severidad": "info",
        "factor_caracteres": 1.1,
    },
    "estilos_inconsistentes": {
        "activo": True,
        "severidad": "warning",
        "fila_encabezado": 1,
    },
    "islas_datos": {
        "activo": True,
        "severidad": "info",
        "hueco_minimo": 3,
    },
    "proteccion": {
        "activo": True,
        "severidad": "info",
        "requerida": False,
    },
    # --- Documentos docx ---
    "docx_titulos_estilos": {
        "activo": True,
        "severidad": "warning",
        "tamano_min_pt": 13,
    },
    "docx_fuentes": {
        "activo": True,
        "severidad": "warning",
        "max_fuentes": 2,
    },
    "docx_margenes": {
        "activo": True,
        "severidad": "warning",
        "min_cm": 1.0,
        "max_cm": 5.0,
    },
    "docx_numeracion_manual": {
        "activo": True,
        "severidad": "warning",
    },
    "docx_parrafos_vacios": {
        "activo": True,
        "severidad": "info",
        "max_consecutivos": 2,
    },
    "docx_tablas_sin_estilo": {
        "activo": True,
        "severidad": "warning",
    },
    "docx_encabezados_pie": {
        "activo": True,
        "severidad": "info",
        "requeridos": False,
    },
    "docx_imagenes": {
        "activo": True,
        "severidad": "info",
    },
    # --- Documentos pdf ---
    "pdf_paginas_vacias": {
        "activo": True,
        "severidad": "warning",
    },
    "pdf_paginas_escasas": {
        "activo": True,
        "severidad": "warning",
        "min_caracteres": 20,
    },
    "pdf_sin_capa_texto": {
        "activo": True,
        "severidad": "info",
    },
    "pdf_rotacion": {
        "activo": True,
        "severidad": "warning",
    },
    "pdf_tamano_paginas": {
        "activo": True,
        "severidad": "info",
        "tolerancia_pt": 1.0,
    },
}

# Solo 'General' y '@' (texto) son senales de celdas numericas sin formato
# explicito: '0', '0.00', '#,##0.00' son formatos numericos validos.
FORMATOS_GENERICOS = {"general", "@"}

_PATRON_ERROR = re.compile(r"^#[A-Z¡/0-9!?-]+!?$")


def cargar_reglas(ruta: str | None = None) -> tuple[dict, list[str]]:
    """Reglas activas: defaults + overrides del archivo JSON.

    Devuelve (reglas, errores). Con ruta None (o archivo inexistente/invalido
    con errores) se usan los defaults; los errores se reportan, no se lanzan,
    para que la CLI pueda avisar sin morir.
    """
    errores = []
    if ruta is None:
        return REGLAS_DEFAULT, errores
    if not os.path.exists(ruta):
        return REGLAS_DEFAULT, [f"archivo de reglas inexistente: {ruta}"]
    try:
        with open(ruta, encoding="utf-8") as f:
            overrides = json.load(f)
    except (json.JSONDecodeError, OSError) as exc:
        return REGLAS_DEFAULT, [f"reglas invalidas ({ruta}): {exc}"]
    if not isinstance(overrides, dict):
        return REGLAS_DEFAULT, [f"reglas invalidas ({ruta}): se espera un objeto JSON"]
    reglas = {}
    for nombre, cfg in REGLAS_DEFAULT.items():
        reglas[nombre] = dict(cfg)
        override = overrides.get(nombre)
        if override is None:
            continue
        if not isinstance(override, dict):
            errores.append(f"check '{nombre}': se espera un objeto")
            continue
        for clave, valor in override.items():
            if clave == "severidad" and valor not in SEVERIDADES:
                errores.append(
                    f"check '{nombre}': severidad invalida '{valor}' "
                    f"(validas: {list(SEVERIDADES)})")
                continue
            reglas[nombre][clave] = valor
    desconocidos = set(overrides) - set(REGLAS_DEFAULT)
    for nombre in sorted(desconocidos):
        errores.append(f"check desconocido: '{nombre}'")
    return reglas, errores


def hallazgo(regla: str, cfg: dict, hoja: str, celda: str, mensaje: str,
             detalle: dict | None = None) -> dict:
    """Construye un hallazgo con la severidad definida por la regla."""
    return {"regla": regla, "severidad": cfg["severidad"], "hoja": hoja,
            "celda": celda, "mensaje": mensaje, "detalle": detalle or {}}


# ---------------------------------------------------------------------------
# Apertura y rango de datos
# ---------------------------------------------------------------------------

def abrir_libro(ruta: str):
    """Abre el libro en modo normal (los estilos no se leen en read_only)."""
    from openpyxl import load_workbook

    return load_workbook(ruta, read_only=False, data_only=False)


def rango_datos(hoja) -> tuple[int, int, int, int]:
    """Rango con datos de la hoja: (min_row, min_col, max_row, max_col).

    Ignora filas/columnas completamente vacias de los extremos. La lectura
    usa ws.iter_rows con values_only=False para inspeccionar las celdas.
    """
    min_r = min_c = None
    max_r = max_c = 0
    for fila in hoja.iter_rows():
        for celda in fila:
            if celda.value is None:
                continue
            if min_r is None:
                min_r = celda.row
                min_c = celda.column
            max_r = celda.row
            max_c = max(max_c, celda.column)
    if min_r is None:
        return 0, 0, 0, 0
    return min_r, min_c, max_r, max_c


def _celda_llena(celda) -> bool:
    return celda.value is not None


# ---------------------------------------------------------------------------
# Checks (cada uno: (hoja, cfg, ctx) -> list[hallazgo])
# ---------------------------------------------------------------------------

def check_encabezados(hoja, cfg, ctx):
    """Encabezado legible: negrita/color/bordes segun lo requerido + nombre
    para cada columna con datos."""
    hallazgos = []
    fila = cfg["fila_encabezado"]
    min_r, min_c, max_r, max_c = ctx["rango"]
    for col in range(min_c, max_c + 1):
        celda = hoja.cell(row=fila, column=col)
        if not _celda_llena(celda):
            # Solo si la columna tiene datos reales por debajo del
            # encabezado (una celda mezclada amplia no es una columna).
            if max_r > fila and not any(
                    hoja.cell(row=r, column=col).value is not None
                    for r in range(fila + 1, max_r + 1)):
                continue
            nombre = f"{hoja.cell(row=fila, column=col).coordinate}"
            hallazgos.append(hallazgo(
                "encabezados", cfg, hoja.title, nombre,
                "columna con datos sin nombre de encabezado"))
            continue
        if cfg["negrita_requerida"] and not (celda.font.b or celda.font.bold):
            hallazgos.append(hallazgo(
                "encabezados", cfg, hoja.title, celda.coordinate,
                f"encabezado '{celda.value}' sin negrita"))
        if cfg["color_fondo_requerido"] and (celda.fill.fill_type is None
                                             or not celda.fill.fgColor.rgb):
            hallazgos.append(hallazgo(
                "encabezados", cfg, hoja.title, celda.coordinate,
                f"encabezado '{celda.value}' sin color de fondo"))
        if cfg["bordes_requeridos"]:
            b = celda.border
            estilos = [b.left.style, b.right.style, b.top.style, b.bottom.style]
            if not any(estilos):
                hallazgos.append(hallazgo(
                    "encabezados", cfg, hoja.title, celda.coordinate,
                    f"encabezado '{celda.value}' sin bordes"))
    return hallazgos


def check_bordes(hoja, cfg, ctx):
    """Celdas del rango de datos sin bordes (muestreo acotado por
    max_hallazgos del contexto)."""
    hallazgos = []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return hallazgos
    limite = ctx.get("max_hallazgos", 500)
    for fila in hoja.iter_rows(min_row=min_r, max_row=max_r,
                               min_col=min_c, max_col=max_c):
        for celda in fila:
            if not _celda_llena(celda):
                continue
            b = celda.border
            estilos = [b.left.style, b.right.style, b.top.style, b.bottom.style]
            if not any(estilos):
                hallazgos.append(hallazgo(
                    "bordes", cfg, hoja.title, celda.coordinate,
                    f"celda '{celda.value}' sin bordes"))
                if len(hallazgos) >= limite:
                    hallazgos.append(hallazgo(
                        "bordes", cfg, hoja.title, "",
                        f"hallazgos limitados a {limite} (celdas sin bordes)"))
                    return hallazgos
    return hallazgos


def check_alineacion(hoja, cfg, ctx):
    """Numericos alineados a izquierda (sospecha de texto) y texto a la
    derecha: senal clasica de datos mal tipados o estilos perdidos."""
    hallazgos = []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return hallazgos
    limite = ctx.get("max_hallazgos", 500)
    for fila in hoja.iter_rows(min_row=min_r + 1, max_row=max_r,
                               min_col=min_c, max_col=max_c):
        for celda in fila:
            valor = celda.value
            if valor is None:
                continue
            alineado = (celda.alignment.horizontal or "").lower()
            if isinstance(valor, (int, float)) and not isinstance(valor, bool):
                if alineado == "left":
                    hallazgos.append(hallazgo(
                        "alineacion", cfg, hoja.title, celda.coordinate,
                        f"valor numerico '{valor}' alineado a la izquierda"))
            elif isinstance(valor, str) and alineado == "right":
                hallazgos.append(hallazgo(
                    "alineacion", cfg, hoja.title, celda.coordinate,
                    f"texto '{valor[:30]}' alineado a la derecha"))
            if len(hallazgos) >= limite:
                hallazgos.append(hallazgo(
                    "alineacion", cfg, hoja.title, "",
                    f"hallazgos limitados a {limite} (alineaciones)"))
                return hallazgos
    return hallazgos


def check_anchos(hoja, cfg, ctx):
    """Anchos de columna fuera de [minimo, maximo] (solo columnas con
    datos; las columnas sin ancho explicito usan el default 8.43)."""
    hallazgos = []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return hallazgos
    dims = hoja.column_dimensions
    for col in range(min_c, max_c + 1):
        letra = _indice_a_letra(col)
        ancho = None
        if letra in dims and dims[letra].width is not None:
            ancho = float(dims[letra].width)
        if ancho is None:
            continue  # ancho default de Excel: no reportado como anomalia
        if ancho < cfg["minimo"]:
            hallazgos.append(hallazgo(
                "anchos", cfg, hoja.title, f"columna {letra}",
                f"ancho {ancho:g} < minimo {cfg['minimo']}"))
        elif ancho > cfg["maximo"]:
            hallazgos.append(hallazgo(
                "anchos", cfg, hoja.title, f"columna {letra}",
                f"ancho {ancho:g} > maximo {cfg['maximo']}"))
    return hallazgos


def _es_numerico(valor) -> bool:
    return isinstance(valor, (int, float)) and not isinstance(valor, bool)


def check_formato_numero(hoja, cfg, ctx):
    """Numeros como texto (data_type 's' con valor convertible) y flotantes
    con formato 'General' (sin formato numerico explicito). Los enteros en
    General (anos, conteos) se consideran normales y no se reportan."""
    hallazgos = []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return hallazgos
    limite = ctx.get("max_hallazgos", 500)
    for fila in hoja.iter_rows(min_row=min_r, max_row=max_r,
                               min_col=min_c, max_col=max_c):
        for celda in fila:
            valor = celda.value
            if valor is None:
                continue
            if celda.data_type == "s" and isinstance(valor, str):
                try:
                    float(valor.replace(",", "."))
                except ValueError:
                    continue
                hallazgos.append(hallazgo(
                    "formato_numero", cfg, hoja.title, celda.coordinate,
                    f"numero guardado como texto: '{valor}' (tipado o formato erroneo)"))
                if len(hallazgos) >= limite:
                    hallazgos.append(hallazgo(
                        "formato_numero", cfg, hoja.title, "",
                        f"hallazgos limitados a {limite} (numeros como texto)"))
                    return hallazgos
                continue
            if _es_numerico(valor) and isinstance(valor, float) and valor != int(valor):
                fmt = (celda.number_format or "General").lower()
                if fmt in FORMATOS_GENERICOS:
                    hallazgos.append(hallazgo(
                        "formato_numero", cfg, hoja.title, celda.coordinate,
                        f"valor numerico '{valor}' con formato '{fmt}' "
                        "(no numerico: los decimales no estan asegurados)"))
                    if len(hallazgos) >= limite:
                        hallazgos.append(hallazgo(
                            "formato_numero", cfg, hoja.title, "",
                            f"hallazgos limitados a {limite} (formato General)"))
                        return hallazgos
    return hallazgos


def check_filtros(hoja, cfg, ctx):
    """La hoja de datos debe tener auto-filtro si se exige."""
    if not cfg["requerido"]:
        return []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return []
    if hoja.auto_filter.ref is None:
        return [hallazgo(
            "filtros", cfg, hoja.title, f"A{min_r}:{_indice_a_letra(max_c)}{max_r}",
            "rango de datos sin auto-filtro")]
    return []


def check_celdas_vacias(hoja, cfg, ctx):
    """Celdas vacias dentro del rango de datos (huecos en medio de una
    fila con datos a ambos lados)."""
    hallazgos = []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return hallazgos
    limite = ctx.get("max_hallazgos", 500)
    for fila in hoja.iter_rows(min_row=min_r + 1, max_row=max_r,
                               min_col=min_c, max_col=max_c):
        llenas = [c for c in fila if _celda_llena(c)]
        if not llenas:
            continue
        for i, celda in enumerate(fila):
            if celda.value is not None or i in (0, len(fila) - 1):
                continue
            if fila[i - 1].value is not None and fila[i + 1].value is not None:
                hallazgos.append(hallazgo(
                    "celdas_vacias", cfg, hoja.title, celda.coordinate,
                    "celda vacia entre valores de la misma fila"))
                if len(hallazgos) >= limite:
                    hallazgos.append(hallazgo(
                        "celdas_vacias", cfg, hoja.title, "",
                        f"hallazgos limitados a {limite} (celdas vacias)"))
                    return hallazgos
    return hallazgos


def check_celdas_mezcladas(hoja, cfg, ctx):
    """Celdas fusionadas: informativo (dificultan analisis y ordenamiento)."""
    rangos = sorted(str(r) for r in hoja.merged_cells.ranges)
    limite = ctx.get("max_hallazgos", 500)
    if not rangos:
        return []
    salida = [hallazgo(
        "celdas_mezcladas", cfg, hoja.title, ", ".join(rangos[:5]),
        f"{len(rangos)} rango(s) de celdas fusionadas"
        + ("" if len(rangos) <= 5 else f" (primeros 5: {', '.join(rangos[:5])})"),
        {"total": len(rangos), "rangos": rangos[:limite]})]
    return salida


def check_ocultas(hoja, cfg, ctx, modo: str):
    """Filas o columnas ocultas: riesgo de datos que no se ven."""
    nombre_regla = "filas_ocultas" if modo == "filas" else "columnas_ocultas"
    hallazgos = []
    if modo == "filas":
        ocultas = [r for r, dim in hoja.row_dimensions.items()
                   if dim.hidden]
        detalle = {"filas": sorted(ocultas)[:20], "total": len(ocultas)}
        texto = "fila(s) oculta(s)"
    else:
        ocultas = [letra for letra, dim in hoja.column_dimensions.items()
                   if dim.hidden]
        detalle = {"columnas": ocultas[:20], "total": len(ocultas)}
        texto = "columna(s) oculta(s)"
    if ocultas:
        muestra = ", ".join(map(str, sorted(ocultas)[:10]))
        hallazgos.append(hallazgo(
            nombre_regla, cfg, hoja.title, muestra,
            f"{len(ocultas)} {texto} (datos ocultos de la presentacion)",
            detalle))
    return hallazgos


def check_errores_formula(hoja, cfg, ctx):
    """Celdas cuyo valor cacheado es un error de formula (#DIV/0!, #N/A...)."""
    hallazgos = []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return hallazgos
    limite = ctx.get("max_hallazgos", 500)
    for fila in hoja.iter_rows(min_row=min_r, max_row=max_r,
                               min_col=min_c, max_col=max_c):
        for celda in fila:
            valor = celda.value
            if isinstance(valor, str) and _PATRON_ERROR.match(valor.strip()):
                hallazgos.append(hallazgo(
                    "errores_formula", cfg, hoja.title, celda.coordinate,
                    f"error de formula: '{valor}'"))
                if len(hallazgos) >= limite:
                    hallazgos.append(hallazgo(
                        "errores_formula", cfg, hoja.title, "",
                        f"hallazgos limitados a {limite} (errores)"))
                    return hallazgos
    return hallazgos


def check_duplicados_encabezado(hoja, cfg, ctx):
    """Nombres de columna duplicados (dificultan identificar columnas)."""
    fila = cfg["fila_encabezado"]
    min_r, min_c, max_r, max_c = ctx["rango"]
    vistos: dict[str, list[str]] = {}
    for col in range(min_c, max_c + 1):
        celda = hoja.cell(row=fila, column=col)
        valor = celda.value
        if valor is None:
            continue
        clave = str(valor)
        vistos.setdefault(clave, []).append(celda.coordinate)
    hallazgos = []
    for nombre, celdas in vistos.items():
        if len(celdas) > 1:
            hallazgos.append(hallazgo(
                "duplicados_encabezado", cfg, hoja.title, ", ".join(celdas),
                f"encabezado duplicado '{nombre}' en {len(celdas)} columnas",
                {"nombre": nombre, "celdas": celdas}))
    return hallazgos


def check_texto_desbordado(hoja, cfg, ctx):
    """Heuristica de texto que probablemente no cabe en su celda: texto
    largo sin wrap_text y con la celda vecina a la derecha ocupada.
    Aproximacion honesta: no hay motor de layout, solo estimacion por
    longitud vs ancho de columna (factor_caracteres).
    """
    hallazgos = []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return hallazgos
    limite = ctx.get("max_hallazgos", 500)
    dims = hoja.column_dimensions
    for fila in hoja.iter_rows(min_row=min_r, max_row=max_r,
                               min_col=min_c, max_col=max_c):
        for i, celda in enumerate(fila):
            valor = celda.value
            if not isinstance(valor, str) or not valor:
                continue
            if (celda.alignment.wrap_text or ""):
                continue
            letra = _indice_a_letra(celda.column)
            ancho = None
            if letra in dims and dims[letra].width is not None:
                ancho = float(dims[letra].width)
            if ancho is None:
                continue
            if len(valor) > ancho * cfg["factor_caracteres"]:
                vecino = None
                if i + 1 < len(fila):
                    vecino = fila[i + 1].value
                if vecino is not None:
                    hallazgos.append(hallazgo(
                        "texto_desbordado", cfg, hoja.title, celda.coordinate,
                        f"texto de {len(valor)} caracteres probablemente cortado "
                        f"(ancho columna {ancho:g}, celda vecina ocupada)",
                        {"texto": valor[:60]}))
                    if len(hallazgos) >= limite:
                        hallazgos.append(hallazgo(
                            "texto_desbordado", cfg, hoja.title, "",
                            f"hallazgos limitados a {limite} (textos largos)"))
                        return hallazgos
    return hallazgos


def check_estilos_inconsistentes(hoja, cfg, ctx):
    """Formatos de numero distintos dentro de una misma columna de datos
    (mezcla de 'General', decimales distintos, etc.)."""
    hallazgos = []
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return hallazgos
    fila_enc = cfg.get("fila_encabezado", 1)
    for col in range(min_c, max_c + 1):
        formatos: dict[str, list[str]] = {}
        for fila in hoja.iter_rows(min_row=max(min_r, fila_enc + 1),
                                   max_row=max_r, min_col=col, max_col=col):
            for celda in fila:
                if not _es_numerico(celda.value):
                    continue
                fmt = (celda.number_format or "General").lower()
                formatos.setdefault(fmt, []).append(celda.coordinate)
        if len(formatos) > 1:
            letra = _indice_a_letra(col)
            enc = hoja.cell(row=fila_enc, column=col).value
            hallazgos.append(hallazgo(
                "estilos_inconsistentes", cfg, hoja.title, f"columna {letra}",
                f"{len(formatos)} formatos numericos distintos en la columna "
                f"'{enc}' ({', '.join(sorted(formatos))})",
                {fmt: celdas[:10] for fmt, celdas in formatos.items()}))
    return hallazgos


def check_islas_datos(hoja, cfg, ctx):
    """Filas de datos separadas del cuerpo principal por un hueco >= N filas
    vacias (posibles bloques sin encabezado o tablas superpuestas)."""
    min_r, min_c, max_r, max_c = ctx["rango"]
    if not max_r:
        return []
    hueco = cfg["hueco_minimo"]
    filas_llenas = []
    for fila in hoja.iter_rows(min_row=min_r + 1, max_row=max_r,
                               min_col=min_c, max_col=max_c):
        if any(_celda_llena(c) for c in fila):
            filas_llenas.append(fila[0].row)
    islas = []
    anterior = None
    for r in filas_llenas:
        if anterior is not None and r - anterior - 1 >= hueco:
            islas.append((anterior, r))
        anterior = r
    if not islas:
        return []
    salida = [hallazgo(
        "islas_datos", cfg, hoja.title,
        ", ".join(f"filas {a+1}-{b-1}" for a, b in islas[:5]),
        f"{len(islas)} hueco(s) de >= {hueco} filas vacias dentro del rango de datos",
        {"huecos": [(a + 1, b - 1) for a, b in islas]})]
    return salida


def check_proteccion(hoja, cfg, ctx):
    """Hoja protegida o no (informativo de cumplimiento)."""
    if cfg["requerida"] and not hoja.protection.sheet:
        return [hallazgo(
            "proteccion", cfg, hoja.title, "",
            "hoja sin proteccion (se requiere por configuracion)")]
    if not hoja.protection.sheet:
        return []
    return []


CHECKS = {
    "encabezados": check_encabezados,
    "bordes": check_bordes,
    "alineacion": check_alineacion,
    "anchos": check_anchos,
    "formato_numero": check_formato_numero,
    "filtros": check_filtros,
    "celdas_vacias": check_celdas_vacias,
    "celdas_mezcladas": check_celdas_mezcladas,
    "filas_ocultas": lambda h, c, x: check_ocultas(h, c, x, "filas"),
    "columnas_ocultas": lambda h, c, x: check_ocultas(h, c, x, "columnas"),
    "errores_formula": check_errores_formula,
    "duplicados_encabezado": check_duplicados_encabezado,
    "texto_desbordado": check_texto_desbordado,
    "estilos_inconsistentes": check_estilos_inconsistentes,
    "islas_datos": check_islas_datos,
    "proteccion": check_proteccion,
}


# ---------------------------------------------------------------------------
# Checks de documentos docx (python-docx)
# ---------------------------------------------------------------------------

def check_docx_titulos_estilos(doc, cfg, ctx):
    """Textos con negrita o tamano grande SIN estilo de titulo (heading):
    formato manual en lugar de estilos de documento (mantenibilidad)."""
    hallazgos = []
    limite = ctx.get("max_hallazgos", 500)
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        estilo = (p.style.name or "").lower()
        if estilo.startswith("heading") or estilo.startswith("titulo") \
                or estilo.startswith("title"):
            continue
        for r in p.runs:
            tamano = r.font.size.pt if r.font.size else None
            if r.bold or (tamano and tamano >= cfg["tamano_min_pt"]):
                hallazgos.append(hallazgo(
                    "docx_titulos_estilos", cfg, "documento", "parrafo",
                    f"estilo manual (negrita/tamano {tamano}pt) en vez de "
                    f"estilo de titulo: '{p.text[:60]}'",
                    {"texto": p.text[:200]}))
                if len(hallazgos) >= limite:
                    hallazgos.append(hallazgo(
                        "docx_titulos_estilos", cfg, "documento", "",
                        f"hallazgos limitados a {limite} (estilos manuales)"))
                    return hallazgos
                break
    return hallazgos


def check_docx_fuentes(doc, cfg, ctx):
    """Mas fuentes distintas de las permitidas en el cuerpo del documento."""
    fuentes = {}
    for p in doc.paragraphs:
        for r in p.runs:
            nombre = r.font.name or "default"
            fuentes.setdefault(nombre, 0)
            fuentes[nombre] += 1
    if len(fuentes) > cfg["max_fuentes"]:
        top = ", ".join(f"{n} ({c})" for n, c in
                        sorted(fuentes.items(), key=lambda kv: -kv[1])[:5])
        return [hallazgo(
            "docx_fuentes", cfg, "documento", "cuerpo",
            f"{len(fuentes)} fuentes distintas en el cuerpo (max "
            f"{cfg['max_fuentes']}): {top}",
            {"fuentes": fuentes})]
    return []


def check_docx_margenes(doc, cfg, ctx):
    """Margenes de seccion fuera del rango [min_cm, max_cm]."""
    hallazgos = []
    for i, seccion in enumerate(doc.sections, start=1):
        for nombre, valor in (("izquierdo", seccion.left_margin),
                              ("derecho", seccion.right_margin),
                              ("superior", seccion.top_margin),
                              ("inferior", seccion.bottom_margin)):
            if valor is None:
                continue
            cm = valor.cm
            if cm < cfg["min_cm"] or cm > cfg["max_cm"]:
                hallazgos.append(hallazgo(
                    "docx_margenes", cfg, "documento", f"seccion {i}",
                    f"margen {nombre} de {cm:.2f} cm fuera de "
                    f"[{cfg['min_cm']}, {cfg['max_cm']}] cm"))
    return hallazgos


def check_docx_numeracion_manual(doc, cfg, ctx):
    """Listas numeradas/viñetadas escritas a mano (estilo Normal) en vez de
    estilos de lista de documento."""
    hallazgos = []
    limite = ctx.get("max_hallazgos", 500)
    patron = re.compile(r"^\s*(\d+[.)]|[-*•])\s")
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        estilo = (p.style.name or "").lower()
        if estilo.startswith("list"):
            continue
        if patron.match(p.text):
            hallazgos.append(hallazgo(
                "docx_numeracion_manual", cfg, "documento", "parrafo",
                f"numeracion manual en vez de estilo de lista: '{p.text[:60]}'",
                {"texto": p.text[:200]}))
            if len(hallazgos) >= limite:
                hallazgos.append(hallazgo(
                    "docx_numeracion_manual", cfg, "documento", "",
                    f"hallazgos limitados a {limite} (numeracion manual)"))
                return hallazgos
    return hallazgos


def check_docx_parrafos_vacios(doc, cfg, ctx):
    """Secuencias de parrafos vacios consecutivos (saltos mal formados)."""
    huecos = []
    racha = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            racha += 1
        else:
            if racha > cfg["max_consecutivos"]:
                huecos.append(racha)
            racha = 0
    if racha > cfg["max_consecutivos"]:
        huecos.append(racha)
    if not huecos:
        return []
    return [hallazgo(
        "docx_parrafos_vacios", cfg, "documento", "cuerpo",
        f"{len(huecos)} secuencia(s) de > {cfg['max_consecutivos']} parrafos "
        f"vacios consecutivos (max racha {max(huecos)})",
        {"rachas": huecos})]


def check_docx_tablas_sin_estilo(doc, cfg, ctx):
    """Tablas sin estilo de tabla con bordes visibles (el default
    'Normal Table' no dibuja bordes)."""
    hallazgos = []
    for i, tabla in enumerate(doc.tables, start=1):
        nombre = (tabla.style.name if tabla.style else "") or ""
        sin_bordes = (not nombre.lower() or
                      nombre.lower() in ("normal table", "tabla normal") or
                      "grid" not in nombre.lower())
        if sin_bordes:
            hallazgos.append(hallazgo(
                "docx_tablas_sin_estilo", cfg, "documento", f"tabla {i}",
                f"tabla {i} sin estilo de tabla con bordes visibles "
                f"(estilo: '{nombre or 'ninguno'}')"))
    return hallazgos


def check_docx_encabezados_pie(doc, cfg, ctx):
    """Documento sin encabezado ni pie de pagina (informativo; requerido
    opcional por configuracion)."""
    con_encabezado = any(
        "".join(p.text for p in s.header.paragraphs).strip() for s in doc.sections)
    con_pie = any(
        "".join(p.text for p in s.footer.paragraphs).strip() for s in doc.sections)
    if cfg["requeridos"] and not (con_encabezado or con_pie):
        return [hallazgo(
            "docx_encabezados_pie", cfg, "documento", "",
            "documento sin encabezado ni pie de pagina (requeridos)")]
    if not (con_encabezado or con_pie):
        return [hallazgo(
            "docx_encabezados_pie", cfg, "documento", "",
            "documento sin encabezado ni pie de pagina (identificacion del "
            "documento en pagina)")]
    return []


def check_docx_imagenes(doc, cfg, ctx):
    """Conteo de imagenes insertadas (informativo)."""
    n = len(doc.inline_shapes)
    if not n:
        return []
    return [hallazgo(
        "docx_imagenes", cfg, "documento", "",
        f"{n} imagen(es) insertada(s) en el documento",
        {"imagenes": n})]


CHECKS_DOCX = {
    "docx_titulos_estilos": check_docx_titulos_estilos,
    "docx_fuentes": check_docx_fuentes,
    "docx_margenes": check_docx_margenes,
    "docx_numeracion_manual": check_docx_numeracion_manual,
    "docx_parrafos_vacios": check_docx_parrafos_vacios,
    "docx_tablas_sin_estilo": check_docx_tablas_sin_estilo,
    "docx_encabezados_pie": check_docx_encabezados_pie,
    "docx_imagenes": check_docx_imagenes,
}


# ---------------------------------------------------------------------------
# Checks de documentos pdf (pypdfium2)
# ---------------------------------------------------------------------------

def check_pdf_paginas_vacias(paginas, cfg, ctx):
    """Paginas sin ningun texto (posible pagina en blanco o render roto)."""
    vacias = [p["indice"] + 1 for p in paginas if not p["texto"].strip()]
    if not vacias:
        return []
    return [hallazgo(
        "pdf_paginas_vacias", cfg, "documento", ", ".join(map(str, vacias[:10])),
        f"{len(vacias)} pagina(s) sin texto (en blanco)",
        {"paginas": vacias[:20]})]


def check_pdf_paginas_escasas(paginas, cfg, ctx):
    """Paginas con muy poco texto: posible contenido cortado o pagina rota."""
    minimo = cfg["min_caracteres"]
    escasas = [(p["indice"] + 1, len(p["texto"].strip()))
               for p in paginas if 0 < len(p["texto"].strip()) < minimo]
    if not escasas:
        return []
    return [hallazgo(
        "pdf_paginas_escasas", cfg, "documento",
        ", ".join(f"pag. {n}" for n, _ in escasas[:10]),
        f"{len(escasas)} pagina(s) con menos de {minimo} caracteres "
        f"(posible contenido cortado)",
        {"paginas": [{"pagina": n, "caracteres": c} for n, c in escasas[:20]]})]


def check_pdf_sin_capa_texto(paginas, cfg, ctx):
    """PDF sin capa de texto (escaneado): los checks de texto no aplican;
    la lectura requiere vision/OCR."""
    total = sum(len(p["texto"].strip()) for p in paginas)
    if total:
        return []
    return [hallazgo(
        "pdf_sin_capa_texto", cfg, "documento", "",
        "PDF sin capa de texto (escaneado o solo imagenes): los checks de "
        "texto no aplican; usar --vision o OCR para el contenido",
        {"paginas": len(paginas)})]


def check_pdf_rotacion(paginas, cfg, ctx):
    """Paginas rotadas respecto a la orientacion natural."""
    rotadas = [p["indice"] + 1 for p in paginas if p["rotacion"]]
    if not rotadas:
        return []
    return [hallazgo(
        "pdf_rotacion", cfg, "documento", ", ".join(map(str, rotadas[:10])),
        f"{len(rotadas)} pagina(s) rotadas (presentacion incorrecta)",
        {"paginas": rotadas[:20]})]


def check_pdf_tamano_paginas(paginas, cfg, ctx):
    """Tamano de pagina distinto entre paginas del mismo documento."""
    tolerancia = cfg["tolerancia_pt"]
    tamanos = {}
    for p in paginas:
        clave = (round(p["ancho"], 1), round(p["alto"], 1))
        tamanos.setdefault(clave, []).append(p["indice"] + 1)
    if len(tamanos) <= 1:
        return []
    detalle = {f"{w:g}x{h:g}pt": paginas_idx[:10]
               for (w, h), paginas_idx in tamanos.items()}
    return [hallazgo(
        "pdf_tamano_paginas", cfg, "documento",
        "; ".join(list(detalle)[:4]),
        f"{len(tamanos)} tamanos de pagina distintos (tolerancia "
        f"{tolerancia} pt)",
        detalle)]


CHECKS_PDF = {
    "pdf_paginas_vacias": check_pdf_paginas_vacias,
    "pdf_paginas_escasas": check_pdf_paginas_escasas,
    "pdf_sin_capa_texto": check_pdf_sin_capa_texto,
    "pdf_rotacion": check_pdf_rotacion,
    "pdf_tamano_paginas": check_pdf_tamano_paginas,
}


# ---------------------------------------------------------------------------
# Revisión de una planilla
# ---------------------------------------------------------------------------

def _indice_a_letra(indice: int) -> str:
    """1 -> A, 27 -> AA (get_column_letter con fallback para hoja.cell)."""
    from openpyxl.utils import get_column_letter

    return get_column_letter(indice)


def revisar_planilla(ruta: str, reglas: dict | None = None,
                     hoja_solo: str | None = None,
                     max_hallazgos: int = 500) -> dict:
    """Revisa el formato y la presentacion de una planilla xlsx.

    Devuelve dict con hojas, hallazgos por regla y resumen por severidad.
    Lanza ValueError para archivos invalidos; KeyError nunca (estructura
    de reglas saneada por cargar_reglas).
    """
    if reglas is None:
        reglas, _ = cargar_reglas()
    libro = abrir_libro(ruta)
    if hoja_solo and hoja_solo not in libro.sheetnames:
        raise ValueError(
            f"hoja '{hoja_solo}' inexistente (hojas: {libro.sheetnames})")

    hojas_a_revisar = libro.sheetnames if not hoja_solo else [hoja_solo]
    hallazgos = []
    for nombre in hojas_a_revisar:
        hoja = libro[nombre]
        ctx = {"rango": rango_datos(hoja), "max_hallazgos": max_hallazgos}
        for regla, fn in CHECKS.items():
            cfg = reglas[regla]
            if not cfg["activo"]:
                continue
            try:
                hallazgos.extend(fn(hoja, cfg, ctx))
            except Exception as exc:  # noqa: BLE001 — un check no debe matar la revision
                hallazgos.append(hallazgo(
                    regla, {**cfg, "severidad": "error"}, nombre, "",
                    f"error interno del check: {exc}"))
    hallazgos = hallazgos[:max_hallazgos]
    resumen = {"error": 0, "warning": 0, "info": 0}
    for h in hallazgos:
        resumen[h["severidad"]] += 1
    return {
        "ok": True,
        "archivo": os.path.basename(ruta),
        "formato": "xlsx",
        "hojas": libro.sheetnames,
        "resumen": resumen,
        "hallazgos": hallazgos,
        "reglas": {"activas": sum(1 for r in reglas.values() if r["activo"]),
                   "desactivadas": [n for n, r in reglas.items() if not r["activo"]]},
    }


# ---------------------------------------------------------------------------
# Comparación entre versiones
# ---------------------------------------------------------------------------

def _estilo_encabezado(celda) -> dict:
    """Resumen comparable del estilo de una celda de encabezado."""
    b = celda.border
    return {
        "negrita": bool(celda.font.b or celda.font.bold),
        "color_fondo": str(celda.fill.fgColor.rgb) if celda.fill and celda.fill.fgColor else None,
        "bordes": [b.left.style, b.right.style, b.top.style, b.bottom.style],
    }


def _valor_normalizado(celda) -> str:
    v = celda.value
    if v is None:
        return ""
    if isinstance(v, float):
        return f"{v:g}"
    return str(v)


def comparar_planillas(ruta_a: str, ruta_b: str, max_diferencias: int = 100) -> dict:
    """Diferencias de estructura, estilos y valores entre dos xlsx."""
    libro_a = abrir_libro(ruta_a)
    libro_b = abrir_libro(ruta_b)
    nombres_a, nombres_b = set(libro_a.sheetnames), set(libro_b.sheetnames)
    diferencias = []

    def _anadir(d: dict):
        diferencias.append(d)

    for nombre in sorted(nombres_a - nombres_b):
        _anadir({"tipo": "hoja", "hoja": nombre, "detalle": "solo en A"})
    for nombre in sorted(nombres_b - nombres_a):
        _anadir({"tipo": "hoja", "hoja": nombre, "detalle": "solo en B"})

    for nombre in sorted(nombres_a & nombres_b):
        a, b = libro_a[nombre], libro_b[nombre]
        ra, rb = rango_datos(a), rango_datos(b)
        if ra != rb:
            _anadir({"tipo": "dimension", "hoja": nombre,
                     "a": list(ra), "b": list(rb),
                     "detalle": f"rango A{ra[0]}:{_indice_a_letra(ra[3])}{ra[2]} vs "
                                f"A{rb[0]}:{_indice_a_letra(rb[3])}{rb[2]}"})
        for col in range(1, max(ra[3], rb[3]) + 1):
            ca = a.cell(row=1, column=col)
            cb = b.cell(row=1, column=col)
            if _valor_normalizado(ca) != _valor_normalizado(cb):
                _anadir({"tipo": "encabezado", "hoja": nombre, "celda": ca.coordinate,
                         "a": _valor_normalizado(ca), "b": _valor_normalizado(cb)})
            elif _estilo_encabezado(ca) != _estilo_encabezado(cb):
                _anadir({"tipo": "estilo_encabezado", "hoja": nombre,
                         "celda": ca.coordinate,
                         "a": _estilo_encabezado(ca), "b": _estilo_encabezado(cb)})
            if len(diferencias) >= max_diferencias:
                break
        if len(diferencias) >= max_diferencias:
            break
        max_filas = max(ra[2], rb[2])
        max_cols = max(ra[3], rb[3])
        for fila_a, fila_b in zip(
                a.iter_rows(min_row=2, max_row=max_filas, max_col=max_cols),
                b.iter_rows(min_row=2, max_row=max_filas, max_col=max_cols)):
            for ca, cb in zip(fila_a, fila_b):
                va, vb = _valor_normalizado(ca), _valor_normalizado(cb)
                if va != vb:
                    _anadir({"tipo": "valor", "hoja": nombre, "celda": ca.coordinate,
                             "a": va, "b": vb})
                    if len(diferencias) >= max_diferencias:
                        break
            if len(diferencias) >= max_diferencias:
                break
        if len(diferencias) >= max_diferencias:
            break

    return {
        "ok": True,
        "archivo_a": os.path.basename(ruta_a),
        "archivo_b": os.path.basename(ruta_b),
        "total_diferencias": len(diferencias),
        "limitado": len(diferencias) >= max_diferencias,
        "diferencias": diferencias,
    }


# ---------------------------------------------------------------------------
# Visión IA 360° (render + VLM local)
# ---------------------------------------------------------------------------

def _soffice_binario() -> str:
    import shutil

    binario = shutil.which("soffice") or shutil.which("libreoffice")
    if not binario:
        raise RuntimeError(
            "se requiere LibreOffice (soffice) en el PATH: conversion y render "
            "de ods/docx/xlsx usan soffice headless")
    return binario


def _convertir_con_soffice(ruta: str, dir_salida: str, formato: str,
                           timeout_s: int = 300) -> str:
    """Convierte un documento con soffice headless. Devuelve la ruta de
    salida (mismo nombre base, extension nueva)."""
    binario = _soffice_binario()
    env = {**os.environ, "TMPDIR": "/var/tmp"}
    cmd = [binario, "--headless", "--convert-to", formato, "--outdir",
           dir_salida, ruta]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True,
                              timeout=timeout_s, env=env)
    except subprocess.TimeoutExpired:
        raise RuntimeError(
            f"soffice supero el timeout de {timeout_s} s convirtiendo {ruta}") from None
    if proc.returncode != 0:
        raise RuntimeError(
            f"soffice fallo (rc={proc.returncode}): {proc.stderr.strip()[-300:]}")
    salida = os.path.join(dir_salida,
                          os.path.splitext(os.path.basename(ruta))[0]
                          + f".{formato}")
    if not os.path.exists(salida):
        candidatos = [f for f in os.listdir(dir_salida) if f.endswith(f".{formato}")]
        if not candidatos:
            raise RuntimeError(
                f"soffice no genero el {formato} esperado para {ruta}")
        salida = os.path.join(dir_salida, candidatos[0])
    return salida


def revisar_ods(ruta: str, reglas: dict | None = None,
                hoja_solo: str | None = None,
                max_hallazgos: int = 500) -> dict:
    """Revisa una planilla ODS normalizandola a xlsx con LibreOffice.

    Los estilos ODS no son XML de Excel: se normalizan via soffice headless
    y los checks de xlsx se aplican sobre el resultado, con una nota de
    transformacion. La integridad de los datos se verifica comparando
    hojas y dimensiones entre el ODS original (odfpy/pandas) y el xlsx
    convertido: si difieren, la revision se marca como no fiable.
    """
    import pandas as pd

    dir_tmp = tempfile.mkdtemp(prefix="revision_ods_")
    try:
        xlsx = _convertir_con_soffice(ruta, dir_tmp, "xlsx")
        integridad = _verificar_integridad_ods(ruta, xlsx)
        resultado = revisar_planilla(xlsx, reglas, hoja_solo, max_hallazgos)
    finally:
        import shutil

        shutil.rmtree(dir_tmp, ignore_errors=True)
    resultado["formato"] = "ods"
    resultado["transformacion"] = {
        "metodo": "soffice headless -> xlsx temporal (estilos normalizados)",
        "integridad": integridad,
    }
    if integridad.get("ok") is False:
        resultado["ok"] = False
    return resultado


def _verificar_integridad_ods(ods: str, xlsx: str) -> dict:
    """Compara hojas y dimensiones entre el ODS original y el xlsx
    convertido (muestreo estructural, no de valores).

    Si odfpy no esta instalado (p. ej. python del sistema sin el venv), la
    integridad queda "no verificada" con aviso: la conversion y los checks
    siguen funcionando, pero sin la garantia de fidelidad.
    """
    try:
        import pandas as pd
    except ImportError as exc:  # pragma: no cover — pandas es dependencia base
        return {"ok": None, "aviso": f"pandas no importable: {exc}"}
    try:
        orig = pd.read_excel(ods, engine="odf", sheet_name=None)
    except ImportError as exc:
        return {"ok": None,
                "aviso": (f"odfpy no instalado: integridad no verificada "
                          f"({exc})")}
    except Exception as exc:  # noqa: BLE001 — cualquier error de lectura
        return {"ok": False,
                "error": f"no se pudo leer el ODS original: {exc}"}
    try:
        conv = pd.read_excel(xlsx, sheet_name=None)
    except Exception as exc:  # noqa: BLE001
        return {"ok": False,
                "error": f"no se pudo leer el xlsx convertido: {exc}"}
    if set(orig) != set(conv):
        return {"ok": False,
                "error": (f"hojas distintas tras la conversion: "
                          f"{sorted(orig)} vs {sorted(conv)}")}
    for hoja in orig:
        o, c = orig[hoja].shape, conv[hoja].shape
        if o != c:
            return {"ok": False,
                    "error": f"hoja '{hoja}': dimensiones {o} vs {c} tras la conversion"}
    return {"ok": True,
            "hojas": len(orig),
            "celdas": sum(len(orig[h]) * len(orig[h].columns) for h in orig)}


def render_libro_a_pngs(ruta: str, dir_salida: str | None = None,
                        timeout_s: int = 300) -> list[str]:
    """Convierte el libro/documento a imagenes PNG (una por pagina del PDF).

    Pipeline: LibreOffice headless -> PDF -> pypdfium2 -> PNG. Requiere
    soffice/libreoffice en el PATH. Devuelve las rutas PNG ordenadas.
    """
    dir_tmp = dir_salida or tempfile.mkdtemp(prefix="revision_render_")
    pdf = _convertir_con_soffice(ruta, dir_tmp, "pdf", timeout_s)

    import pypdfium2 as pdfium

    doc = pdfium.PdfDocument(pdf)
    pngs = []
    for i in range(len(doc)):
        pagina = doc[i]
        imagen = pagina.render(scale=2).to_pil()
        png = os.path.join(dir_tmp, f"pagina_{i + 1}.png")
        imagen.save(png)
        pngs.append(png)
    doc.close()
    return pngs


def render_pdf_a_pngs(ruta: str, dir_salida: str | None = None,
                      escala: float = 2.0) -> list[str]:
    """Render directo de un PDF a PNG con pypdfium2 (sin soffice).

    Para la Vision IA 360 de PDFs: el render nativo es fiel a la
    presentacion final, sin normalizacion intermedia.
    """
    import pypdfium2 as pdfium

    dir_tmp = dir_salida or tempfile.mkdtemp(prefix="revision_render_")
    doc = pdfium.PdfDocument(ruta)
    pngs = []
    for i in range(len(doc)):
        pagina = doc[i]
        imagen = pagina.render(scale=escala).to_pil()
        png = os.path.join(dir_tmp, f"pagina_{i + 1}.png")
        imagen.save(png)
        pngs.append(png)
    doc.close()
    return pngs


PROMPT_RUBRICA = (
    "Evalua el DISENO y la PRESENTACION de esta pagina de una planilla de calculo. "
    "Responde EXACTAMENTE una linea por dimension con el formato "
    "'<dimension>: nota /10 | comentario breve'.\n"
    "Dimensiones: legibilidad (tamano y contraste de texto), coherencia de estilo "
    "(encabezados, bordes, colores consistentes), uso del color (adecuado y "
    "accesible), estructura y alineacion (columnas alineadas, orden claro), "
    "formato de datos (numeros con formato correcto, fechas consistentes), "
    "presentacion general (apariencia profesional para entrega a terceros).\n"
    "Solo notas numericas de 0 a 10, sin decimales."
)


# Dimensiones aceptadas de la rubrica (prefijos): cualquier otra etiqueta
# con nota se marca como respuesta no conforme.
DIMENSIONES_VALIDAS = ("legibilidad", "coherencia", "uso del color",
                       "estructura", "formato de datos", "presentación",
                       "presentacion")

_DIM_PALABRAS = r"([\wáéíóúñ]+(?:\s+[\wáéíóúñ]+)*)"
_PATRON_DIM_PRIMERO = re.compile(
    rf"^\s*{_DIM_PALABRAS}\s*[:|-]\s*(\d{{1,2}})\s*(?:/10)?\s*\|?(.*)$",
    re.IGNORECASE)
_PATRON_NOTA_PRIMERO = re.compile(
    rf"^\s*(\d{{1,2}})\s*(?:/10)?\s*:?\s*{_DIM_PALABRAS}\s*\|?(.*)$",
    re.IGNORECASE)


def _prefijo_canonico(dim: str) -> str | None:
    """Prefijo de rubrica que inicia la dimension (clave canonica).

    'Legibilidad excelente' -> 'legibilidad'; 'coherencia de estilo' ->
    'coherencia'; None si no corresponde a ninguna dimension valida.
    """
    dim = dim.lower().strip()
    for v in sorted(DIMENSIONES_VALIDAS, key=len, reverse=True):
        if dim.startswith(v):
            return v
    return None


def _parsear_rubrica(texto: str) -> tuple[dict, dict, int]:
    """Extrae 'dimension: nota/10 | comentario' o 'nota/10: dimension | ...'.

    Devuelve (notas, comentarios, no_conformes): las claves son los
    prefijos canonicos de la rubrica; no_conformes cuenta lineas con nota
    fuera de rubrica (el VLM a veces inventa sus propias dimensiones).
    """
    notas, comentarios = {}, {}
    no_conformes = 0
    for linea in texto.splitlines():
        m = _PATRON_DIM_PRIMERO.match(linea.strip()) or \
            _PATRON_NOTA_PRIMERO.match(linea.strip())
        if not m:
            continue
        if m.re is _PATRON_DIM_PRIMERO:
            dim, nota_s, coment = m.group(1), m.group(2), m.group(3)
        else:
            nota_s, dim, coment = m.group(1), m.group(2), m.group(3)
        nota = int(nota_s)
        if not (0 <= nota <= 10):
            continue
        canonica = _prefijo_canonico(dim)
        if canonica is None:
            no_conformes += 1
            continue
        notas[canonica] = nota
        comentarios[canonica] = coment.strip()[:200]
    return notas, comentarios, no_conformes


def vision_360(ruta: str, motor: str = "ollama", host: str = "127.0.0.1",
               device: str = "cuda", modelo: str = "gemma3:4b",
               timeout_s: int = 1800) -> dict:
    """Evaluacion visual de diseno/presentacion con un VLM local.

    Renderiza el libro a PNGs y pregunta la rubrica de 6 dimensiones por
    pagina. Reutiliza los motores validados de scripts/bateria_360.py
    (docbee en subproceso aislado; ollama con arranque bajo demanda).
    """
    if motor not in ("docbee", "ollama"):
        return {"ok": False,
                "error": f"motor invalido: {motor} (docbee|ollama)"}
    try:
        pngs = _render_segun_formato(ruta)
    except RuntimeError as exc:
        return {"ok": False, "error": f"render: {exc}"}

    sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "scripts"))
    try:
        from bateria_360 import run_docbee, run_ollama
    except ImportError as exc:
        return {"ok": False, "error": f"scripts/bateria_360.py no importable: {exc}"}

    resultados = []
    for i, png in enumerate(pngs):
        if motor == "docbee":
            res = run_docbee(png, PROMPT_RUBRICA, device, timeout_s=timeout_s)
        else:
            res = run_ollama(png, PROMPT_RUBRICA, host, modelo, timeout_s=timeout_s)
        entrada = {"pagina": i + 1, "imagen": png}
        if not res.get("ok"):
            entrada["error"] = res.get("error")
            entrada["tiempo_s"] = res.get("total_s")
            resultados.append(entrada)
            continue
        texto = res.get("texto", "")
        notas, comentarios, no_conformes = _parsear_rubrica(texto)
        entrada.update({
            "texto": texto[:4000],
            "notas": notas,
            "comentarios": comentarios,
            "no_conformes": no_conformes,
            "tiempo_s": res.get("total_s") or res.get("tiempo_s"),
        })
        resultados.append(entrada)

    return {"ok": True, "motor": motor,
            "modelo": "PP-DocBee-2B" if motor == "docbee" else modelo,
            "paginas": len(pngs), "resultados": resultados}


# ---------------------------------------------------------------------------
# Despacho por formato
# ---------------------------------------------------------------------------

SOPORTADOS = {"xlsx", "xlsm", "ods", "docx", "pdf"}
# Formatos que admiten comparacion entre versiones (--comparar)
COMPARABLES = {"xlsx", "xlsm", "ods"}
# Formatos con hoja/s (para --hoja)
CON_HOJAS = {"xlsx", "xlsm", "ods"}


def _render_segun_formato(ruta: str) -> list[str]:
    """Render a PNGs segun el formato: PDF nativo con pypdfium2 (fiel a la
    presentacion final); el resto via LibreOffice (normalizacion)."""
    ext = os.path.splitext(ruta)[1].lstrip(".").lower()
    if ext == "pdf":
        return render_pdf_a_pngs(ruta)
    return render_libro_a_pngs(ruta)


def revisar_docx(ruta: str, reglas: dict | None = None,
                 max_hallazgos: int = 500) -> dict:
    """Revisa el formato y presentacion de un documento docx (python-docx).

    Checks de estilo de documento: titulos con estilos en vez de negrita
    manual, consistencia de fuentes, margenes, numeracion manual, parrafos
    vacios, tablas sin estilo, encabezados/pie e imagenes.
    """
    from docx import Document

    if reglas is None:
        reglas, _ = cargar_reglas()
    doc = Document(ruta)
    ctx = {"doc": doc, "max_hallazgos": max_hallazgos}
    hallazgos = []
    for regla, fn in CHECKS_DOCX.items():
        cfg = reglas[regla]
        if not cfg["activo"]:
            continue
        try:
            hallazgos.extend(fn(doc, cfg, ctx))
        except Exception as exc:  # noqa: BLE001 — un check no debe matar la revision
            hallazgos.append(hallazgo(
                regla, {**cfg, "severidad": "error"}, "documento", "",
                f"error interno del check: {exc}"))
    hallazgos = hallazgos[:max_hallazgos]
    resumen = {"error": 0, "warning": 0, "info": 0}
    for h in hallazgos:
        resumen[h["severidad"]] += 1
    return {
        "ok": True,
        "archivo": os.path.basename(ruta),
        "formato": "docx",
        "estructura": {"secciones": len(doc.sections),
                       "parrafos": len(doc.paragraphs),
                       "tablas": len(doc.tables)},
        "resumen": resumen,
        "hallazgos": hallazgos,
        "reglas": {"activas": sum(1 for r in reglas.values() if r["activo"]),
                   "desactivadas": [n for n, r in reglas.items() if not r["activo"]]},
    }


def revisar_pdf(ruta: str, reglas: dict | None = None,
                max_hallazgos: int = 500) -> dict:
    """Revisa la presentacion de un PDF con pypdfium2.

    Checks de presentacion: paginas vacias, paginas con muy poco texto,
    ausencia de capa de texto (escaneado), rotacion y tamanos de pagina
    distintos. El contenido visual se evalua con --vision (render nativo).
    """
    import pypdfium2 as pdfium

    if reglas is None:
        reglas, _ = cargar_reglas()
    doc = pdfium.PdfDocument(ruta)
    try:
        paginas = []
        for i in range(len(doc)):
            pagina = doc[i]
            tp = pagina.get_textpage()
            ancho, alto = pagina.get_size()
            paginas.append({
                "indice": i,
                "texto": (tp.get_text_range() or ""),
                "rotacion": pagina.get_rotation(),
                "ancho": ancho,
                "alto": alto,
            })
    finally:
        doc.close()
    ctx = {"paginas": paginas, "max_hallazgos": max_hallazgos}
    hallazgos = []
    for regla, fn in CHECKS_PDF.items():
        cfg = reglas[regla]
        if not cfg["activo"]:
            continue
        try:
            hallazgos.extend(fn(paginas, cfg, ctx))
        except Exception as exc:  # noqa: BLE001 — un check no debe matar la revision
            hallazgos.append(hallazgo(
                regla, {**cfg, "severidad": "error"}, "documento", "",
                f"error interno del check: {exc}"))
    hallazgos = hallazgos[:max_hallazgos]
    resumen = {"error": 0, "warning": 0, "info": 0}
    for h in hallazgos:
        resumen[h["severidad"]] += 1
    return {
        "ok": True,
        "archivo": os.path.basename(ruta),
        "formato": "pdf",
        "estructura": {"paginas": len(paginas),
                       "caracteres": sum(len(p["texto"].strip()) for p in paginas)},
        "resumen": resumen,
        "hallazgos": hallazgos,
        "reglas": {"activas": sum(1 for r in reglas.values() if r["activo"]),
                   "desactivadas": [n for n, r in reglas.items() if not r["activo"]]},
    }


def revisar_documento(ruta: str, reglas: dict | None = None,
                      comparar: str | None = None,
                      hoja_solo: str | None = None,
                      max_hallazgos: int = 500,
                      vision: str | None = None,
                      vision_host: str = "127.0.0.1",
                      vision_device: str = "cuda",
                      vision_modelo: str = "gemma3:4b") -> dict:
    """Revision completa de un documento (despacho por extension)."""
    if not os.path.exists(ruta):
        return {"ok": False, "error": f"el archivo no existe: {ruta}"}
    ext = os.path.splitext(ruta)[1].lstrip(".").lower()
    if ext not in SOPORTADOS:
        return {"ok": False,
                "error": f"formato '.{ext}' no soportado (soportados: "
                         f"{sorted(SOPORTADOS)})"}
    if hoja_solo and ext not in CON_HOJAS:
        return {"ok": False,
                "error": f"--hoja solo aplica a {sorted(CON_HOJAS)}, no a '{ext}'"}
    if comparar and ext not in COMPARABLES:
        return {"ok": False,
                "error": f"--comparar solo aplica a {sorted(COMPARABLES)}, no a '{ext}'"}
    try:
        if ext in ("xlsx", "xlsm"):
            resultado = revisar_planilla(ruta, reglas, hoja_solo, max_hallazgos)
        elif ext == "ods":
            resultado = revisar_ods(ruta, reglas, hoja_solo, max_hallazgos)
        elif ext == "docx":
            resultado = revisar_docx(ruta, reglas, max_hallazgos)
        else:
            resultado = revisar_pdf(ruta, reglas, max_hallazgos)
    except ValueError as exc:
        return {"ok": False, "error": str(exc)}
    except Exception as exc:  # noqa: BLE001 — archivos corruptos u otros
        return {"ok": False, "error": f"no se pudo revisar: {exc}"}
    if comparar:
        if not os.path.exists(comparar):
            return {"ok": False, "error": f"el archivo a comparar no existe: {comparar}"}
        # openpyxl no lee ods: ambos lados se normalizan a xlsx temporal
        # antes de comparar (solo si hacen falta).
        ruta_a, dir_a = ruta, None
        ruta_b, dir_b = comparar, None
        import shutil

        try:
            if ext == "ods":
                dir_a = tempfile.mkdtemp(prefix="revision_comp_")
                ruta_a = _convertir_con_soffice(ruta, dir_a, "xlsx")
            ext_b = os.path.splitext(comparar)[1].lstrip(".").lower()
            if ext_b == "ods":
                dir_b = tempfile.mkdtemp(prefix="revision_comp_")
                ruta_b = _convertir_con_soffice(comparar, dir_b, "xlsx")
            resultado["comparacion"] = comparar_planillas(ruta_a, ruta_b)
        finally:
            for d in (dir_a, dir_b):
                if d:
                    shutil.rmtree(d, ignore_errors=True)
    if vision:
        resultado["vision"] = vision_360(ruta, motor=vision, host=vision_host,
                                         device=vision_device, modelo=vision_modelo)
    return resultado


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Revision de formato y presentacion de planillas (xlsx) y documentos")
    parser.add_argument("archivo", help="Ruta a la planilla (xlsx/xlsm)")
    parser.add_argument("--reglas", help="Archivo JSON de reglas (opcional, usa defaults)")
    parser.add_argument("--comparar", help="Segunda planilla para comparar versiones")
    parser.add_argument("--hoja", help="Revisar solo esta hoja")
    parser.add_argument("--max-hallazgos", type=int, default=500,
                        help="Limite de hallazgos (default: 500)")
    parser.add_argument("--vision", choices=["docbee", "ollama"],
                        help="Evaluar diseno/presentacion con un VLM local (Vision IA 360)")
    parser.add_argument("--modelo", default="gemma3:4b",
                        help="Modelo ollama para --vision (default: gemma3:4b)")
    parser.add_argument("--host", default="127.0.0.1", help="Host de ollama")
    parser.add_argument("--device", default="cuda", help="Device para docbee")
    parser.add_argument("--salida", choices=["json", "md"], default="json",
                        help="Formato del reporte (default: json)")
    args = parser.parse_args()

    reglas, errores_reglas = cargar_reglas(args.reglas)
    t0 = time.monotonic()
    resultado = revisar_documento(
        args.archivo, reglas=reglas, comparar=args.comparar, hoja_solo=args.hoja,
        max_hallazgos=args.max_hallazgos, vision=args.vision,
        vision_host=args.host, vision_device=args.device, vision_modelo=args.modelo)
    resultado["tiempo_s"] = round(time.monotonic() - t0, 1)
    if errores_reglas:
        resultado["errores_reglas"] = errores_reglas
    if args.salida == "md":
        print(_resumen_markdown(resultado))
    else:
        print(json.dumps(resultado, ensure_ascii=False, indent=2))
    if not resultado.get("ok", True):
        sys.exit(2)


def _resumen_markdown(resultado: dict) -> str:
    """Resumen legible del reporte (tabla markdown, sin dependencias)."""
    lineas = [f"# Revision: {resultado.get('archivo', '?')}",
              f"- formato: {resultado.get('formato', '?')}",
              f"- tiempo: {resultado.get('tiempo_s', '?')} s"]
    resumen = resultado.get("resumen")
    if resumen:
        lineas.append(f"- hallazgos: {resumen.get('error', 0)} error, "
                      f"{resumen.get('warning', 0)} warning, "
                      f"{resumen.get('info', 0)} info")
    lineas.append("")
    hallazgos = resultado.get("hallazgos", [])
    if not hallazgos:
        lineas.append("Sin hallazgos.")
    else:
        lineas.append("| Regla | Severidad | Hoja | Celda | Mensaje |")
        lineas.append("|---|---|---|---|---|")
        for h in hallazgos:
            msg = str(h.get("mensaje", "")).replace("|", "\\|")
            lineas.append(f"| {h.get('regla', '')} | {h.get('severidad', '')} "
                          f"| {h.get('hoja', '')} | {h.get('celda', '')} | {msg} |")
    comparacion = resultado.get("comparacion")
    if comparacion:
        lineas.append("")
        lineas.append(f"## Comparacion: {comparacion.get('total_diferencias', 0)} "
                      f"diferencias (limitado: {comparacion.get('limitado', False)})")
    vision = resultado.get("vision")
    if vision:
        lineas.append("")
        lineas.append(f"## Vision IA 360 ({vision.get('motor')})")
        for res in vision.get("resultados", []):
            lineas.append(f"- pagina {res.get('pagina')}: "
                          f"{res.get('notas') or res.get('error')}")
    return "\n".join(lineas)


if __name__ == "__main__":
    main()
