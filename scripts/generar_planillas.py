#!/usr/bin/env python3
"""Genera documentos de prueba sinteticos para el modulo de revision
(xlsx, ods, docx y pdf; datos propios, sin contenido de terceros).

xlsx:
  correcta.xlsx    : cumple las reglas por defecto de revision.py
  con_fallos.xlsx  : viola la mayoria de los checks (referencia negativa)
  v1.xlsx/v2.xlsx  : dos versiones del mismo informe para --comparar
ods:
  correcta.ods     : misma tabla que correcta.xlsx (conversion normalizada)
docx:
  documento_correcto.docx : estilos de documento (titulos, fuente unica)
  documento_con_fallos.docx : formato manual (negrita, fuentes, numeracion)
pdf:
  documento.pdf    : PDF generado (docx correcto convertido con soffice)

Uso:
  python3 scripts/generar_planillas.py [--salida DIR]
"""

import argparse
import os
import sys

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


def _estilo_cabecera(hoja, ncols, fila=1):
    """Estilo de encabezado estandar (negrita + fondo + bordes)."""
    for col in range(1, ncols + 1):
        celda = hoja.cell(row=fila, column=col)
        celda.font = Font(bold=True, color="FFFFFF")
        celda.fill = PatternFill("solid", fgColor="4472C4")
        borde = Side(style="thin", color="000000")
        celda.border = Border(left=borde, right=borde, top=borde, bottom=borde)
        celda.alignment = Alignment(horizontal="center")


def _borde_fino(celda):
    borde = Side(style="thin", color="000000")
    celda.border = Border(left=borde, right=borde, top=borde, bottom=borde)


def _anchos(hoja, dims):
    for letra, ancho in dims.items():
        hoja.column_dimensions[letra].width = ancho


def generar_correcta(ruta):
    """Planilla que cumple los defaults: sin hallazgos esperados."""
    wb = Workbook()
    hoja = wb.active
    hoja.title = "Ventas"
    cabeceras = ["Producto", "Region", "Anio", "Unidades", "Importe"]
    hoja.append(cabeceras)
    _estilo_cabecera(hoja, len(cabeceras))
    filas = [
        ("Bicicleta urbana", "Norte", 2024, 120, 45200.5),
        ("Bicicleta urbana", "Sur", 2024, 85, 32100.0),
        ("Bicicleta de montana", "Norte", 2024, 60, 28400.75),
        ("Bicicleta de montana", "Sur", 2024, 42, 19900.25),
        ("Bicicleta urbana", "Norte", 2025, 138, 52800.0),
        ("Bicicleta urbana", "Sur", 2025, 91, 34900.5),
        ("Bicicleta de montana", "Norte", 2025, 74, 35100.0),
        ("Bicicleta de montana", "Sur", 2025, 50, 23700.75),
    ]
    for i, fila in enumerate(filas, start=2):
        for j, valor in enumerate(fila, start=1):
            celda = hoja.cell(row=i, column=j, value=valor)
            _borde_fino(celda)
            if j in (4, 5):
                celda.number_format = "0.00" if j == 5 else "0"
    _anchos(hoja, {"A": 22, "B": 12, "C": 8, "D": 10, "E": 12})
    hoja.auto_filter.ref = f"A1:E{len(filas) + 1}"
    wb.save(ruta)


def generar_con_fallos(ruta):
    """Planilla que viola los checks por defecto (referencia negativa)."""
    wb = Workbook()
    hoja = wb.active
    hoja.title = "Datos"
    hoja.append(["Producto", "Region", "Anio", "Unidades", "Importe", "Notas"])
    # sin negrita ni fondo ni bordes en el encabezado: check encabezados
    # ancho de A estrecho (texto desbordado con vecino ocupado)
    _anchos(hoja, {"A": 6, "B": 10, "C": 10, "D": 10, "E": 10, "F": 30})
    filas = [
        ("Bicicleta urbana de gama alta con accesorios", "Norte", 2024, "120", 45200.5, "#DIV/0!"),
        ("Bicicleta urbana", None, 2024, 85, 32100.0, "bien"),
        ("Bicicleta de montana", "Norte", 2024, 60, 28400.75, "bien"),
        (None, None, None, None, None, None),          # hueco de 4 filas (islas)
        (None, None, None, None, None, None),
        (None, None, None, None, None, None),
        (None, None, None, None, None, None),
        ("Bicicleta de montana", "Sur", 2024, 42, 19900.25, "bien"),
        ("Bicicleta urbana", "Norte", 2025, 138, 52800.0, "bien"),
        ("Bicicleta urbana", "Sur", 2025, 91, 34900.5, "bien"),
    ]
    for i, fila in enumerate(filas, start=2):
        for j, valor in enumerate(fila, start=1):
            if valor is None:
                continue
            hoja.cell(row=i, column=j, value=valor)
    # formato numerico inconsistente: D general, E mezcla General/"0.00"
    for i in range(2, 10):
        d = hoja.cell(row=i, column=4)
        if d.value is not None:
            d.number_format = "General"
        e = hoja.cell(row=i, column=5)
        if e.value is not None:
            e.number_format = "0.00" if i % 2 == 0 else "General"
    # encabezado duplicado (columna "Notas" repetida en G) y datos fuera
    hoja["G1"] = "Notas"
    hoja["G2"] = "duplicado"
    # celda mezclada y celda vacia entre valores (fila 3, columna B)
    hoja.merge_cells("I1:J2")
    hoja["I1"] = "Zona libre"
    hoja.row_dimensions[2].hidden = True   # fila oculta
    hoja.column_dimensions["F"].hidden = True  # columna oculta
    wb.save(ruta)


def generar_versiones(dir_salida):
    """v1.xlsx y v2.xlsx: diferencias para --comparar."""
    wb1 = Workbook()
    hoja = wb1.active
    hoja.title = "Informe"
    hoja.append(["Concepto", "Importe"])
    _estilo_cabecera(hoja, 2)
    hoja.append(["Ingresos", 10000.0])
    hoja.append(["Gastos", 7000.0])
    hoja.append(["Resultado", 3000.0])
    for fila in hoja.iter_rows(min_row=2, max_row=4):
        for celda in fila:
            _borde_fino(celda)
            if celda.column == 2:
                celda.number_format = "0.00"
    _anchos(hoja, {"A": 14, "B": 12})
    wb1.save(os.path.join(dir_salida, "v1.xlsx"))

    wb2 = Workbook()
    hoja = wb2.active
    hoja.title = "Informe"
    hoja.append(["Concepto", "Importe Final"])   # encabezado cambiado
    _estilo_cabecera(hoja, 2)
    hoja.append(["Ingresos", 10500.0])           # valor cambiado
    hoja.append(["Gastos", 7000.0])
    hoja.append(["Resultado", 3500.0])           # valor cambiado
    hoja.append(["Impuestos", 500.0])            # fila nueva
    for fila in hoja.iter_rows(min_row=2, max_row=5):
        for celda in fila:
            _borde_fino(celda)
            if celda.column == 2:
                celda.number_format = "0.00"
    _anchos(hoja, {"A": 14, "B": 12})
    wb2.save(os.path.join(dir_salida, "v2.xlsx"))


def generar_docx(ruta_correcto, ruta_fallos):
    """Documento correcto (estilos de documento) y con fallos (formato
    manual) para los checks docx."""
    from docx import Document
    from docx.shared import Cm, Pt

    def base_documento():
        doc = Document()
        for seccion in doc.sections:
            seccion.left_margin = Cm(2.5)
            seccion.right_margin = Cm(2.5)
            seccion.top_margin = Cm(2.5)
            seccion.bottom_margin = Cm(2.5)
        return doc

    # --- correcto: estilos de titulo, una sola fuente, lista con estilo ---
    doc = base_documento()
    doc.add_heading("Informe trimestral de atencion", level=1)
    doc.add_heading("Resumen ejecutivo", level=2)
    doc.add_paragraph(
        "Durante el trimestre se atendieron 1.240 llamadas, un 8% mas que "
        "el periodo anterior. La satisfaccion media fue de 4,6 sobre 5.")
    doc.add_heading("Detalle por servicio", level=2)
    for item in ("Atencion domiciliaria: 640", "Teleasistencia: 410",
                 "Acompañamiento: 190"):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Tabla resumen", level=2)
    tabla = doc.add_table(rows=3, cols=2)
    tabla.style = "Table Grid"
    for i, (servicio, llamadas) in enumerate(
            (("Domiciliaria", 640), ("Teleasistencia", 410))):
        tabla.cell(i + 1, 0).text = servicio
        tabla.cell(i + 1, 1).text = str(llamadas)
    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            if run.font.name is None:
                run.font.name = "Calibri"
    doc.save(ruta_correcto)

    # --- con fallos: negrita manual, fuentes mezcladas, margenes minimos,
    # numeracion manual, parrafos vacios y tabla sin estilo ---
    doc = base_documento()
    for seccion in doc.sections:
        seccion.left_margin = Cm(0.4)
        seccion.right_margin = Cm(0.4)
    p = doc.add_paragraph()
    run = p.add_run("TITULO ESCRITO A MANO")
    run.bold = True
    run.font.size = Pt(18)
    p2 = doc.add_paragraph()
    r1 = p2.add_run("Texto en Calibri ")
    r1.font.name = "Calibri"
    r2 = p2.add_run("mezclado con Times New Roman")
    r2.font.name = "Times New Roman"
    doc.add_paragraph("1. Primer paso del proceso")
    doc.add_paragraph("2. Segundo paso del proceso")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("3. Tercer paso del proceso")
    tabla = doc.add_table(rows=2, cols=2)  # sin estilo: bordes invisibles
    tabla.cell(0, 0).text = "Concepto"
    tabla.cell(0, 1).text = "Valor"
    tabla.cell(1, 0).text = "Total"
    tabla.cell(1, 1).text = "123"
    doc.save(ruta_fallos)


def generar_ods(ruta, xlsx_origen):
    """Misma planilla que correcta.xlsx en formato ODS, con los estilos
    conservados (soffice normaliza xlsx -> ods; pandas to_excel no aplica
    estilos y daria una planilla 'sin formato' de partida)."""
    import shutil
    import subprocess

    binario = shutil.which("soffice") or shutil.which("libreoffice")
    if not binario:
        print("[AVISO] soffice no esta en el PATH: no se genera el ODS")
        return False
    dir_salida = os.path.dirname(ruta)
    env = {**os.environ, "TMPDIR": "/var/tmp"}
    proc = subprocess.run(
        [binario, "--headless", "--convert-to", "ods", "--outdir", dir_salida,
         xlsx_origen], capture_output=True, text=True, timeout=300, env=env)
    if proc.returncode != 0:
        print(f"[AVISO] soffice fallo al generar el ODS: {proc.stderr[-200:]}")
        return False
    ods_soffice = os.path.join(
        dir_salida, os.path.splitext(os.path.basename(xlsx_origen))[0] + ".ods")
    if ods_soffice != ruta and os.path.exists(ods_soffice):
        os.rename(ods_soffice, ruta)
    return True


def generar_pdf(ruta, docx_origen):
    """Convierte el docx correcto a PDF con soffice (si esta disponible)."""
    import shutil
    import subprocess

    binario = shutil.which("soffice") or shutil.which("libreoffice")
    if not binario:
        print("[AVISO] soffice no esta en el PATH: no se genera el PDF")
        return False
    dir_salida = os.path.dirname(ruta)
    env = {**os.environ, "TMPDIR": "/var/tmp"}
    proc = subprocess.run(
        [binario, "--headless", "--convert-to", "pdf", "--outdir", dir_salida,
         docx_origen], capture_output=True, text=True, timeout=300, env=env)
    if proc.returncode != 0:
        print(f"[AVISO] soffice fallo al generar el PDF: {proc.stderr[-200:]}")
        return False
    pdf_soffice = os.path.join(
        dir_salida, os.path.splitext(os.path.basename(docx_origen))[0] + ".pdf")
    if pdf_soffice != ruta and os.path.exists(pdf_soffice):
        os.rename(pdf_soffice, ruta)
    return True


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--salida", default="ejemplos/planillas",
                        help="Directorio de salida (default: ejemplos/planillas)")
    args = parser.parse_args()

    os.makedirs(args.salida, exist_ok=True)
    generar_correcta(os.path.join(args.salida, "correcta.xlsx"))
    generar_con_fallos(os.path.join(args.salida, "con_fallos.xlsx"))
    generar_versiones(args.salida)
    ok_ods = generar_ods(os.path.join(args.salida, "correcta.ods"),
                         os.path.join(args.salida, "correcta.xlsx"))
    generar_docx(os.path.join(args.salida, "documento_correcto.docx"),
                 os.path.join(args.salida, "documento_con_fallos.docx"))
    ok_pdf = generar_pdf(os.path.join(args.salida, "documento.pdf"),
                         os.path.join(args.salida, "documento_correcto.docx"))
    print(f"[OK] documentos generados en {args.salida}: "
          "correcta.xlsx, con_fallos.xlsx, v1.xlsx, v2.xlsx"
          + (", correcta.ods" if ok_ods else " (sin ods)")
          + ", documento_correcto.docx, documento_con_fallos.docx"
          + (", documento.pdf" if ok_pdf else " (sin pdf)"))


if __name__ == "__main__":
    main()
